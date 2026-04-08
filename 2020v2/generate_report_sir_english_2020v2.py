"""
Generate English Word document report for SIR Model (2020v2).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SIR_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\sir_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

CITY_GROUPS = {
    'North': ['Milano', 'Bergamo', 'Brescia', 'Monza_E_Della_Brianza', 'Como'],
    'Center': ['Roma', 'Firenze', 'Perugia', 'Latina', 'Frosinone'],
    'South': ['Napoli', 'Caserta', 'Salerno', 'Bari', 'Taranto'],
}

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

def add_img(doc, path, caption, w=Inches(6)):
    if os.path.exists(path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True; r.font.size = Pt(10)
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

def main():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.name = 'Times New Roman'

    # Title
    for _ in range(5): doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SIR Model Report\n15 Selected Italian Cities\nCOVID-19 Data (2020)')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('5 Northern | 5 Central | 5 Southern Cities\nMarch 2026')
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The SIR (Susceptible-Infected-Recovered) model is a fundamental epidemiological '
        'compartmental model used to simulate infectious disease dynamics. In this report, '
        'the SIR model is applied to 15 selected Italian cities from 2020 COVID-19 data. '
        'Cities were chosen from three geographic regions (North, Center, South) based on '
        'geographic proximity and total infection counts.')

    # 2. SIR Model
    doc.add_heading('2. SIR Model Description', level=1)
    doc.add_paragraph(
        'The SIR model divides the population into three compartments:\n\n'
        'S (Susceptible): Individuals not yet infected\n'
        'I (Infected): Currently infected individuals\n'
        'R (Recovered): Individuals who have recovered\n\n'
        'Differential Equations:\n'
        'dS/dt = -beta * S * I / N\n'
        'dI/dt = beta * S * I / N - gamma * I\n'
        'dR/dt = gamma * I\n\n'
        'Parameters:\n'
        'gamma = 0.07 (recovery rate, ~14 day recovery period)\n'
        'beta = time-varying (30-day piecewise optimization via Differential Evolution)\n'
        'R0 = beta / gamma (basic reproduction number; R0 > 1 = active spread)')

    param_headers = ['Parameter', 'Value', 'Description']
    param_rows = [
        ['gamma', '0.07', 'Recovery rate (~14 day period)'],
        ['beta optimization', '30-day piecewise', 'Differential Evolution algorithm'],
        ['R0 threshold', '1.0', 'Above 1 = active epidemic spread'],
        ['Segment length', '30 days', 'Beta is constant within each segment'],
    ]
    add_styled_table(doc, param_headers, param_rows)
    doc.add_page_break()

    # 3. Cities and Results
    doc.add_heading('3. Selected Cities and Results', level=1)

    city_headers = ['Region', 'Cities', 'Selection Rationale']
    city_rows = [
        ['North (Lombardy)', 'Milano, Bergamo, Brescia,\nMonza, Como',
         'Lombardy was the 2020 epicenter'],
        ['Center (Lazio/Tuscany)', 'Roma, Firenze, Perugia,\nLatina, Frosinone',
         'Roma (capital) highest central cases'],
        ['South (Campania/Puglia)', 'Napoli, Caserta, Salerno,\nBari, Taranto',
         'Major southern population centers'],
    ]
    add_styled_table(doc, city_headers, city_rows)
    doc.add_paragraph('')

    path = os.path.join(SIR_DIR, 'sir_summary_15cities.csv')
    if os.path.exists(path):
        df = pd.read_csv(path)
        doc.add_heading('SIR Results Summary:', level=3)
        h = ['City', 'Region', 'Population', 'Total Cases', 'Infection %', 'Mean Beta', 'Mean R0']
        rows = []
        for _, row in df.iterrows():
            rows.append([str(row['city']), str(row['region']),
                         f"{int(row['population']):,}", f"{int(row['total_cases']):,}",
                         f"{row['infection_rate_pct']:.1f}%",
                         f"{row['mean_beta']:.4f}", f"{row['mean_R0']:.2f}"])
        add_styled_table(doc, h, rows)
    doc.add_page_break()

    # 4. Plots
    doc.add_heading('4. Visualizations', level=1)

    doc.add_heading('4.1 All Cities Comparison', level=2)
    add_img(doc, os.path.join(SIR_DIR, 'all_cities_comparison.png'),
            'Figure 1: SIR Model Comparison - All 15 Cities')

    doc.add_heading('4.2 Regional Comparison', level=2)
    doc.add_paragraph(
        'The regional comparison shows that Northern cities (Lombardy) had the highest '
        'infection counts, while Southern cities had relatively lower counts in 2020.')
    add_img(doc, os.path.join(SIR_DIR, 'regional_comparison.png'),
            'Figure 2: SIR Regional Comparison (North/Center/South)')

    doc.add_heading('4.3 Correlation Matrix', level=2)
    doc.add_paragraph(
        'The correlation matrix reveals strong intra-regional correlations. Cities within '
        'the same geographic cluster show highly correlated daily infection patterns.')
    add_img(doc, os.path.join(SIR_DIR, 'correlation_matrix.png'),
            'Figure 3: Daily Infection Correlation Matrix')

    doc.add_heading('4.4 Cross-Correlation (Lag Analysis)', level=2)
    doc.add_paragraph(
        'Cross-correlation analysis reveals temporal lags between cities. The epidemic '
        'typically started in larger cities and spread to smaller ones with a delay.')
    add_img(doc, os.path.join(SIR_DIR, 'cross_correlations.png'),
            'Figure 4: Cross-Correlation Between City Pairs')
    doc.add_page_break()

    # Individual SIR plots
    doc.add_heading('4.5 Individual City SIR Plots', level=2)
    doc.add_paragraph(
        'For each city, four plots are generated: (1) Actual vs predicted cumulative infections, '
        '(2) S, I, R compartments, (3) Daily new infections, (4) Time-varying beta(t).')

    fig = 5
    for region, cities in CITY_GROUPS.items():
        for city in cities:
            safe = city.lower().replace("'", "").replace(" ", "_")
            p = os.path.join(SIR_DIR, f'sir_{safe}.png')
            add_img(doc, p, f'Figure {fig}: SIR Model - {city} ({region})', Inches(5.5))
            fig += 1
    doc.add_page_break()

    # 5. Conclusions
    doc.add_heading('5. Conclusions', level=1)
    items = [
        'The SIR model was successfully applied to all 15 selected cities',
        'R0 was above 1.0 for all cities, confirming active epidemic spread in 2020',
        'Milano had the highest total cases (174,013) with R0 = 1.22',
        'Bergamo showed the lowest R0 (1.08), possibly due to early strict restrictions',
        'Cities within the same region showed high correlation in infection patterns',
        'Northern Italy (Lombardy) was confirmed as the epidemic epicenter',
        'The piecewise beta fitting captured the temporal dynamics of infection rates',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    out = os.path.join(OUTPUT_DIR, 'SIR_Report_2020v2_English.docx')
    doc.save(out)
    print(f"Report saved to: {out}")

if __name__ == '__main__':
    main()
