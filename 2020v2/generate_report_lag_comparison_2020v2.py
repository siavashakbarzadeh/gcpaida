"""
Generate comparison report: CGP Lag=7 vs Lag=15 (English + Farsi)
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

LAG7_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\cgp_results"
LAG15_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\cgp_results_lag15"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

def add_img(doc, path, caption, w=Inches(6)):
    if os.path.exists(path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True; r.font.size = Pt(10)
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

def generate_english():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.name = 'Times New Roman'

    # Title
    for _ in range(4): doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('CGP Lag Parameter Study\nLag=7 vs Lag=15 Comparison\n15 Italian Cities (2020)')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('How does increasing temporal lookback from 7 to 15 days\naffect CGP prediction quality?\n\nMarch 2026')
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Study Objective', level=1)
    doc.add_paragraph(
        'This report compares two CGP configurations that differ only in the lag parameter:\n\n'
        '- Lag = 7: Uses the previous 7 days of infection data as features\n'
        '- Lag = 15: Uses the previous 15 days of infection data as features\n\n'
        'All other parameters remain identical:\n'
        '  - Same 15 cities, same addition-only functions\n'
        '  - Same CGP architecture (3x8 grid, levels-back=8)\n'
        '  - Same R-squared threshold (0.9)\n'
        '  - Same lockdown analysis periods\n\n'
        'Key Question: Does a longer temporal lookback (15 days) improve prediction '
        'quality by capturing more of the virus incubation period, or does it introduce '
        'noise from irrelevant older data?')

    doc.add_page_break()

    # 2. R2 Comparison
    doc.add_heading('2. R-squared Comparison', level=1)

    comp_path = os.path.join(LAG15_DIR, 'lag7_vs_lag15_comparison.csv')
    if os.path.exists(comp_path):
        df = pd.read_csv(comp_path)
        headers = ['City', 'Region', 'R2 (Lag=7)', 'R2 (Lag=15)', 'Delta', 'Improved?']
        rows = []
        for _, row in df.iterrows():
            delta = row['r2_delta']
            delta_str = f"+{delta:.4f}" if delta > 0 else f"{delta:.4f}"
            rows.append([
                str(row['city']), str(row['region']),
                f"{row['r2_lag7']:.4f}", f"{row['r2_lag15']:.4f}",
                delta_str, str(row['r2_improved'])
            ])
        add_table(doc, headers, rows)

        # Summary stats
        n_improved = len(df[df['r2_improved'] == 'Yes'])
        n_total = len(df)
        avg_delta = df['r2_delta'].mean()
        sig7 = len(df[df['r2_lag7'] > 0.9])
        sig15 = len(df[df['r2_lag15'] > 0.9])

        doc.add_paragraph('')
        doc.add_paragraph(
            f'Summary Statistics:\n'
            f'- Cities where R2 improved with lag=15: {n_improved}/{n_total}\n'
            f'- Average R2 change: {avg_delta:+.4f}\n'
            f'- Significant connections (R2>0.9) with lag=7: {sig7}\n'
            f'- Significant connections (R2>0.9) with lag=15: {sig15}')

    doc.add_paragraph('')
    add_img(doc, os.path.join(LAG15_DIR, 'lag7_vs_lag15_comparison.png'),
            'Figure 1: R2 and Active Inputs Comparison - Lag=7 vs Lag=15')

    doc.add_page_break()

    # 3. Lockdown comparison
    doc.add_heading('3. Lockdown Analysis Comparison', level=1)
    doc.add_paragraph(
        'The lockdown analysis was repeated with lag=15 to see how the increased '
        'temporal context affects the three-period comparison (pre-lockdown, during '
        'lockdown, after easing).')

    add_img(doc, os.path.join(LAG15_DIR, 'lockdown_lag7_vs_lag15.png'),
            'Figure 2: Lockdown R2 Comparison - Lag=7 vs Lag=15 (3 periods)')

    # Lockdown table
    ld7_path = os.path.join(LAG7_DIR, 'lockdown_comparison.csv')
    ld15_path = os.path.join(LAG15_DIR, 'lockdown_comparison.csv')
    if os.path.exists(ld7_path) and os.path.exists(ld15_path):
        ld7 = pd.read_csv(ld7_path)
        ld15 = pd.read_csv(ld15_path)
        headers = ['City', 'During(7)', 'During(15)', 'After(7)', 'After(15)']
        rows = []
        for _, r7 in ld7.iterrows():
            city = r7['city']
            r15_row = ld15[ld15['city'] == city]
            if len(r15_row) > 0:
                r15 = r15_row.iloc[0]
                rows.append([city,
                    f"{r7.get('during_lockdown_inter_r2', 0):.4f}",
                    f"{r15.get('during_lockdown_inter_r2', 0):.4f}",
                    f"{r7.get('after_easing_inter_r2', 0):.4f}",
                    f"{r15.get('after_easing_inter_r2', 0):.4f}"])
        add_table(doc, headers, rows)

    doc.add_page_break()

    # 4. Lag=15 specific plots
    doc.add_heading('4. Lag=15 Detailed Results', level=1)
    add_img(doc, os.path.join(LAG15_DIR, 'connection_heatmap.png'),
            'Figure 3: Connection Heatmap (Lag=15)')
    add_img(doc, os.path.join(LAG15_DIR, 'network_graph.png'),
            'Figure 4: Network Graph (Lag=15, R2>0.9)')
    add_img(doc, os.path.join(LAG15_DIR, 'pre_post_lockdown_comparison.png'),
            'Figure 5: Lockdown Period Comparison (Lag=15)')
    add_img(doc, os.path.join(LAG15_DIR, 'lag_vs_r2_combined.png'),
            'Figure 6: R2 vs Lag Sweep (from lag=15 analysis)')

    doc.add_page_break()

    # 5. Conclusions
    doc.add_heading('5. Conclusions', level=1)
    conclusions = [
        'Increasing lag from 7 to 15 generally does NOT improve R2 significantly',
        'Lag=7 produced more significant connections (R2>0.9) than lag=15',
        'With lag=15, the number of input features doubles (15 cities x 15 lags = 225 inputs vs 105)',
        'More input features can cause CGP to struggle finding optimal solutions in 500 generations',
        'The during-lockdown period shows lower R2 with lag=15, likely because the lockdown period is only 70 days and lag=15 consumes more data',
        'The lag sweep analysis confirms optimal lags are typically 7-20 days for most cities',
        'Recommendation: Lag=7 is the better default; lag=15 increases dimensionality without proportional benefit',
    ]
    for item in conclusions:
        doc.add_paragraph(item, style='List Bullet')

    out = os.path.join(OUTPUT_DIR, 'CGP_Lag_Comparison_English.docx')
    doc.save(out)
    print(f"Report saved to: {out}")

def generate_farsi():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.name = 'Arial'

    # Title
    for _ in range(4): doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('مطالعه پارامتر لگ CGP\nمقایسه لگ=۷ و لگ=۱۵\n۱۵ شهر ایتالیا (۲۰۲۰)')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('آیا افزایش بازه زمانی از ۷ به ۱۵ روز\nکیفیت پیش‌بینی CGP را بهبود می‌دهد؟\n\nمارس ۲۰۲۶')
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # 1
    doc.add_heading('۱. هدف مطالعه', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(
        'این گزارش دو پیکربندی CGP را مقایسه می‌کند:\n\n'
        '• لگ = ۷: استفاده از ۷ روز قبل (تحلیل اصلی)\n'
        '• لگ = ۱۵: استفاده از ۱۵ روز قبل (تحلیل جدید)\n\n'
        'سوال کلیدی: آیا افزایش بازه زمانی به ۱۵ روز (دوره نهفتگی ویروس) '
        'کیفیت پیش‌بینی را بهبود می‌دهد یا نویز ناخواسته اضافه می‌کند؟')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # 2
    doc.add_heading('۲. مقایسه R²', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    comp_path = os.path.join(LAG15_DIR, 'lag7_vs_lag15_comparison.csv')
    if os.path.exists(comp_path):
        df = pd.read_csv(comp_path)
        headers = ['شهر', 'منطقه', 'R² (لگ=۷)', 'R² (لگ=۱۵)', 'تغییر', 'بهبود؟']
        rows = []
        for _, row in df.iterrows():
            d = row['r2_delta']
            rows.append([str(row['city']), str(row['region']),
                f"{row['r2_lag7']:.4f}", f"{row['r2_lag15']:.4f}",
                f"+{d:.4f}" if d > 0 else f"{d:.4f}",
                'بله' if row['r2_improved'] == 'Yes' else 'خیر'])
        add_table(doc, headers, rows)

        n_imp = len(df[df['r2_improved'] == 'Yes'])
        doc.add_paragraph('')
        p = doc.add_paragraph(
            f'خلاصه: از {len(df)} شهر، {n_imp} شهر با لگ=۱۵ بهبود یافت.\n'
            f'میانگین تغییر R²: {df["r2_delta"].mean():+.4f}')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_img(doc, os.path.join(LAG15_DIR, 'lag7_vs_lag15_comparison.png'),
            'شکل ۱: مقایسه R² و ورودی‌های فعال - لگ=۷ در مقابل لگ=۱۵')
    doc.add_page_break()

    # 3
    doc.add_heading('۳. مقایسه لاک‌دون', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_img(doc, os.path.join(LAG15_DIR, 'lockdown_lag7_vs_lag15.png'),
            'شکل ۲: مقایسه لاک‌دون - لگ=۷ و لگ=۱۵')

    # 4
    doc.add_heading('۴. نمودارهای لگ=۱۵', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_img(doc, os.path.join(LAG15_DIR, 'connection_heatmap.png'), 'شکل ۳: نقشه حرارتی (لگ=۱۵)')
    add_img(doc, os.path.join(LAG15_DIR, 'network_graph.png'), 'شکل ۴: گراف شبکه (لگ=۱۵)')
    add_img(doc, os.path.join(LAG15_DIR, 'pre_post_lockdown_comparison.png'), 'شکل ۵: لاک‌دون (لگ=۱۵)')
    doc.add_page_break()

    # 5
    doc.add_heading('۵. نتیجه‌گیری', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    items = [
        'افزایش لگ از ۷ به ۱۵ عموماً R² را بهبود نمی‌دهد',
        'لگ=۷ ارتباطات معنادار بیشتری (R²>0.9) پیدا کرد',
        'لگ=۱۵ تعداد ورودی‌ها را دو برابر می‌کند (۲۲۵ به جای ۱۰۵)',
        'ورودی‌های بیشتر باعث می‌شود CGP در ۵۰۰ نسل به بهینه نرسد',
        'در دوره لاک‌دون، لگ=۱۵ عملکرد ضعیف‌تری دارد (فقط ۷۰ روز داده)',
        'پیشنهاد: لگ=۷ پیش‌فرض بهتری است',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = os.path.join(OUTPUT_DIR, 'CGP_Lag_Comparison_Farsi.docx')
    doc.save(out)
    print(f"Report saved to: {out}")

if __name__ == '__main__':
    generate_english()
    generate_farsi()
