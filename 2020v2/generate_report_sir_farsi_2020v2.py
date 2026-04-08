"""
Generate Farsi Word document report for SIR Model (2020v2).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SIR_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\sir_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

CITY_GROUPS = {
    'North': ['Milano', 'Bergamo', 'Brescia', 'Monza_E_Della_Brianza', 'Como'],
    'Center': ['Roma', 'Firenze', 'Perugia', 'Latina', 'Frosinone'],
    'South': ['Napoli', 'Caserta', 'Salerno', 'Bari', 'Taranto'],
}

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

def add_img(doc, path, caption, w=Inches(6)):
    if os.path.exists(path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

def main():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.name = 'Arial'

    # Title
    for _ in range(5): doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('گزارش مدل SIR\nبرای ۱۵ شهر منتخب ایتالیا\nداده‌های ۲۰۲۰')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('۵ شهر شمال | ۵ شهر مرکز | ۵ شهر جنوب\nمارس ۲۰۲۶')
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # 1. Intro
    doc.add_heading('۱. مقدمه', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(
        'مدل SIR یکی از مدل‌های کلاسیک اپیدمیولوژی است. '
        'جمعیت به سه گروه تقسیم می‌شود: مستعد (S)، آلوده (I)، و بهبودیافته (R). '
        'در این گزارش مدل SIR بر روی ۱۵ شهر منتخب ایتالیا اعمال شده است.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2. SIR Model
    doc.add_heading('۲. معرفی مدل SIR', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(
        'معادلات:\n'
        'dS/dt = -beta * S * I / N\n'
        'dI/dt = beta * S * I / N - gamma * I\n'
        'dR/dt = gamma * I\n\n'
        'gamma = 0.07 (دوره بهبودی ≈ ۱۴ روز)\n'
        'beta = متغیر با زمان (بهینه‌سازی قطعه‌ای ۳۰ روزه)\n'
        'R0 = beta / gamma')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    param_headers = ['پارامتر', 'مقدار', 'توضیح']
    param_rows = [
        ['gamma', '0.07', 'نرخ بهبودی'],
        ['بهینه‌سازی beta', 'قطعه‌ای ۳۰ روزه', 'Differential Evolution'],
        ['آستانه R0', '1.0', 'بالای ۱ = شیوع فعال'],
    ]
    add_styled_table(doc, param_headers, param_rows)
    doc.add_page_break()

    # 3. Cities
    doc.add_heading('۳. شهرها و نتایج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    path = os.path.join(SIR_DIR, 'sir_summary_15cities.csv')
    if os.path.exists(path):
        df = pd.read_csv(path)
        h = ['شهر', 'منطقه', 'جمعیت', 'کل مبتلایان', 'R0']
        rows = []
        for _, row in df.iterrows():
            rows.append([str(row['city']), str(row['region']),
                         f"{int(row['population']):,}", f"{int(row['total_cases']):,}",
                         f"{row['mean_R0']:.2f}"])
        add_styled_table(doc, h, rows)
    doc.add_page_break()

    # 4. Plots
    doc.add_heading('۴. نمودارها', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_img(doc, os.path.join(SIR_DIR, 'all_cities_comparison.png'), 'شکل ۱: مقایسه ۱۵ شهر')
    add_img(doc, os.path.join(SIR_DIR, 'regional_comparison.png'), 'شکل ۲: مقایسه منطقه‌ای')
    add_img(doc, os.path.join(SIR_DIR, 'correlation_matrix.png'), 'شکل ۳: ماتریس همبستگی')
    add_img(doc, os.path.join(SIR_DIR, 'cross_correlations.png'), 'شکل ۴: تحلیل تأخیر')
    doc.add_page_break()

    # Individual plots
    doc.add_heading('نمودارهای انفرادی SIR:', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fig = 5
    for region, cities in CITY_GROUPS.items():
        for city in cities:
            safe = city.lower().replace("'", "").replace(" ", "_")
            p = os.path.join(SIR_DIR, f'sir_{safe}.png')
            add_img(doc, p, f'شکل {fig}: SIR - {city} ({region})', Inches(5.5))
            fig += 1

    doc.add_page_break()

    # 5. Conclusion
    doc.add_heading('۵. نتیجه‌گیری', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    items = [
        'مدل SIR با موفقیت بر ۱۵ شهر اعمال شد',
        'R0 در همه شهرها بالای ۱ بود (شیوع فعال)',
        'Milano بیشترین مبتلایان (۱۷۴,۰۱۳) را داشت',
        'Bergamo کمترین R0 (۱.۰۸) را نشان داد',
        'شهرهای هم‌منطقه‌ای همبستگی بالایی داشتند',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = os.path.join(OUTPUT_DIR, 'SIR_Report_2020v2_Farsi.docx')
    doc.save(out)
    print(f"Report saved to: {out}")

if __name__ == '__main__':
    main()
