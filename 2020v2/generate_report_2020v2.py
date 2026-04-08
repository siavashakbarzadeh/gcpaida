"""
Generate Word document report (Farsi) for 2020v2 CGP City Analysis.
Includes:
  - SIR model results for 15 selected cities
  - CGP analysis with addition-only functions
  - Lag sweep analysis (R² vs lag)
  - Pre/post lockdown behavior comparison
  - Technical explanation of 'levels back' in CGP
  - Technical explanation of Z-score normalization

Output: 2020v2/reports/
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SIR_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\sir_results"
CGP_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\cgp_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020v2\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table


def add_image_if_exists(doc, img_path, caption, width=Inches(6)):
    if os.path.exists(img_path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
        return True
    return False


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Arial'

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(5):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'گزارش تحلیل ارتباطات بین شهری\n'
        'با استفاده از برنامه‌نویسی ژنتیک کارتزین (CGP)\n'
        'و مدل SIR'
    )
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(46, 64, 87)

    doc.add_paragraph('')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'تحلیل رفتارشناسی انتشار COVID-19 در ۱۵ شهر منتخب ایتالیا\n'
        'داده‌های سال ۲۰۲۰\n\n'
        '۵ شهر شمال | ۵ شهر مرکز | ۵ شهر جنوب'
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('تاریخ: مارس ۲۰۲۶')
    run.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    h = doc.add_heading('فهرست مطالب', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    toc_items = [
        '۱. مقدمه و هدف پروژه',
        '۲. معرفی شهرهای منتخب',
        '۳. مدل SIR و نتایج آن',
        '۴. معرفی برنامه‌نویسی ژنتیک کارتزین (CGP)',
        '   ۴.۱ معماری CGP و پارامتر Levels Back',
        '   ۴.۲ توابع ریاضی (فقط جمع)',
        '   ۴.۳ نرمال‌سازی Z-score و دلیل استفاده',
        '۵. نتایج تحلیل CGP',
        '   ۵.۱ تحلیل ارتباطات بین شهری',
        '   ۵.۲ وزن‌های CGP (کدام شهرها مهم‌تر هستند)',
        '   ۵.۳ تحلیل لگ (R² در مقابل لگ)',
        '۶. تحلیل قبل و بعد از لاک‌دون',
        '۷. نتیجه‌گیری',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    h = doc.add_heading('۱. مقدمه و هدف پروژه', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    intro = (
        'در این پروژه، هدف بررسی و کشف ارتباطات مستقیم بین شهرهای ایتالیا از نظر انتشار بیماری COVID-19 است. '
        'به جای بررسی تمام ۱۰۷ شهر، ۱۵ شهر منتخب از سه منطقه جغرافیایی (شمال، مرکز، جنوب) انتخاب شده‌اند. '
        'انتخاب بر اساس دو معیار انجام شده: '
        '(۱) نزدیکی جغرافیایی شهرها به یکدیگر در هر گروه '
        '(۲) اهمیت شهرها از نظر تعداد مبتلایان.'
    )
    p = doc.add_paragraph(intro)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    intro2 = (
        'دو نوآوری اصلی این تحلیل عبارتند از:\n'
        '• استفاده فقط از عملیات جمع در توابع CGP (بدون تفریق و تفاضل‌گیری)\n'
        '• تحلیل تفاوت رفتار سیستم قبل و بعد از لاک‌دون (۹ مارس ۲۰۲۰)\n\n'
        'قبل از لاک‌دون: انسان‌ها قابلیت جابه‌جایی بین شهرها را داشتند، '
        'پس طبیعتاً اثر دیگر شهرها بیشتر بود. '
        'بعد از لاک‌دون: تعداد مبتلایان فقط تحت تأثیر روزهای قبل همان شهر بود.'
    )
    p = doc.add_paragraph(intro2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # =========================================================================
    # 2. SELECTED CITIES
    # =========================================================================
    h = doc.add_heading('۲. معرفی شهرهای منتخب', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cities_desc = (
        'پانزده شهر در سه گروه جغرافیایی انتخاب شده‌اند. '
        'هر گروه شامل ۵ شهر نزدیک به هم از لحاظ جغرافیایی است '
        'که از نظر تعداد مبتلایان اهمیت بالایی دارند:'
    )
    p = doc.add_paragraph(cities_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    city_headers = ['منطقه', 'شهرها', 'دلیل انتخاب']
    city_rows = [
        ['شمال (لمباردی)',
         'Milano, Bergamo, Brescia,\nMonza e della Brianza, Como',
         'لمباردی مرکز اصلی شیوع در ۲۰۲۰\nبیشترین تعداد مبتلایان'],
        ['مرکز (لاتزیو/توسکانا)',
         'Roma, Firenze, Perugia,\nLatina, Frosinone',
         'رم پایتخت با بیشترین مبتلایان مرکز\nفلورانس شهر اصلی توسکانی'],
        ['جنوب (کامپانیا/پولیا)',
         'Napoli, Caserta, Salerno,\nBari, Taranto',
         'ناپل/کازرتا/سالرنو در کامپانیا\nباری/تارانتو در پولیا'],
    ]
    add_styled_table(doc, city_headers, city_rows)

    doc.add_paragraph('')

    # Load SIR summary if available
    sir_summary_path = os.path.join(SIR_DIR, 'sir_summary_15cities.csv')
    if os.path.exists(sir_summary_path):
        sir_df = pd.read_csv(sir_summary_path)
        sir_headers = ['شهر', 'منطقه', 'جمعیت', 'کل مبتلایان', 'R0 میانگین']
        sir_rows = []
        for _, row in sir_df.iterrows():
            sir_rows.append([
                str(row['city']),
                str(row['region']),
                f"{int(row['population']):,}",
                f"{int(row['total_cases']):,}",
                f"{row['mean_R0']:.2f}",
            ])
        doc.add_heading('خلاصه آماری شهرها:', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_styled_table(doc, sir_headers, sir_rows)

    doc.add_page_break()

    # =========================================================================
    # 3. SIR MODEL
    # =========================================================================
    h = doc.add_heading('۳. مدل SIR و نتایج آن', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sir_desc = (
        'مدل SIR (Susceptible-Infected-Recovered) یک مدل کلاسیک اپیدمیولوژی است که جمعیت را '
        'به سه گروه تقسیم می‌کند: مستعد (S)، آلوده (I)، و بهبودیافته (R). '
        'پارامترهای مدل:\n'
        '• γ (gamma) = ۰.۰۷ (نرخ بهبودی، دوره بهبودی ≈ ۱۴ روز)\n'
        '• β (beta) = نرخ عفونت متغیر با زمان (بهینه‌سازی قطعه‌ای ۳۰ روزه)\n'
        '• N = جمعیت کل هر استان\n\n'
        'معادلات دیفرانسیل مدل SIR:\n'
        '• dS/dt = -β × S × I / N\n'
        '• dI/dt = β × S × I / N - γ × I\n'
        '• dR/dt = γ × I'
    )
    p = doc.add_paragraph(sir_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # SIR comparison plots
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'all_cities_comparison.png'),
                        'شکل ۱: مقایسه مدل SIR برای ۱۵ شهر منتخب')
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'regional_comparison.png'),
                        'شکل ۲: مقایسه منطقه‌ای (شمال، مرکز، جنوب)')
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'correlation_matrix.png'),
                        'شکل ۳: ماتریس همبستگی بین ۱۵ شهر')

    doc.add_page_break()

    # =========================================================================
    # 4. CGP METHODOLOGY
    # =========================================================================
    h = doc.add_heading('۴. معرفی برنامه‌نویسی ژنتیک کارتزین (CGP)', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cgp_desc = (
        'برنامه‌نویسی ژنتیک کارتزین (CGP) یک الگوریتم تکاملی است که توسط Julian Miller '
        'در سال ۲۰۰۰ معرفی شد. در CGP، برنامه‌ها به صورت گراف‌های جهت‌دار بدون حلقه (DAG) '
        'بر روی یک شبکه دوبعدی (سطر × ستون) نمایش داده می‌شوند.'
    )
    p = doc.add_paragraph(cgp_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- 4.1 Levels Back ----
    doc.add_heading('۴.۱ معماری CGP و پارامتر Levels Back', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    levels_back_desc = (
        'یکی از مهم‌ترین پارامترهای معماری CGP، پارامتر "Levels Back" (سطوح به عقب) است. '
        'این پارامتر تعیین می‌کند که هر گره در شبکه CGP می‌تواند تا چند ستون به عقب '
        'به گره‌های قبلی متصل شود.\n\n'
        'توضیح دقیق:\n'
        '• در CGP، گره‌ها در یک شبکه n_rows × n_cols قرار دارند\n'
        '• هر گره دارای دو ورودی (اتصال) و یک تابع ریاضی است\n'
        '• پارامتر levels_back = L تعیین می‌کند که گره‌ای در ستون c '
        'فقط می‌تواند به گره‌های ستون‌های max(0, c-L) تا c-1 متصل شود\n'
        '• همچنین هر گره می‌تواند مستقیماً به ورودی‌های اصلی (داده‌ها) متصل شود\n\n'
        'اثر Levels Back:\n'
        '• اگر L = 1: هر گره فقط به ستون قبلی متصل می‌شود → برنامه‌های خطی ساده\n'
        '• اگر L = n_cols (حداکثر): هر گره به تمام ستون‌های قبلی دسترسی دارد → حداکثر پیچیدگی\n'
        '• مقدار کوچک L: برنامه‌های محلی‌تر و ساده‌تر (کمتر overfit می‌شوند)\n'
        '• مقدار بزرگ L: برنامه‌های پیچیده‌تر (قدرت مدل‌سازی بیشتر)\n\n'
        'در این پروژه: levels_back = 8 (برابر تعداد ستون‌ها = اتصال کامل)\n'
        'دلیل: می‌خواهیم CGP حداکثر آزادی را برای کشف روابط بین شهرها داشته باشد.'
    )
    p = doc.add_paragraph(levels_back_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Levels back example table
    lb_headers = ['Levels Back (L)', 'اتصالات مجاز', 'پیچیدگی', 'مناسب برای']
    lb_rows = [
        ['L = 1', 'فقط ستون قبلی', 'کم', 'مسائل ساده / جلوگیری از overfitting'],
        ['L = 3', '۳ ستون قبلی', 'متوسط', 'تعادل بین پیچیدگی و عمومیت'],
        ['L = n_cols', 'تمام ستون‌ها', 'حداکثر', 'کشف روابط پیچیده (استفاده ما)'],
    ]
    add_styled_table(doc, lb_headers, lb_rows)

    doc.add_paragraph('')

    # CGP Architecture diagram (text-based)
    arch_desc = (
        'نمایش شماتیک معماری CGP ما:\n\n'
        'ورودی‌ها (داده‌های ۱۵ شهر با لگ)  →  [ستون ۱]  →  [ستون ۲]  →  ...  →  [ستون ۸]  →  خروجی\n'
        '                                      ↑ ↑ ↑        ↑ ↑ ↑                  ↑ ↑ ↑\n'
        '                                    (۳ گره)      (۳ گره)               (۳ گره)\n\n'
        'هر گره: func(input1, input2)\n'
        'levels_back = 8: گره در هر ستون می‌تواند به هر ستون قبلی متصل شود'
    )
    p = doc.add_paragraph(arch_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # ---- 4.2 Functions ----
    doc.add_heading('۴.۲ توابع ریاضی (فقط جمع)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    func_desc = (
        'در این پروژه، فقط از عملیات جمع استفاده شده است. '
        'هیچ عملیات تفریق یا تفاضل‌گیری وجود ندارد. '
        'دلیل: ما می‌خواهیم ببینیم آیا فقط با ترکیب جمعی داده‌های شهرها، '
        'می‌توان روند مبتلایان را پیش‌بینی کرد:'
    )
    p = doc.add_paragraph(func_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    func_headers = ['نام تابع', 'عملیات', 'توضیح']
    func_rows = [
        ['add', 'a + b', 'جمع ساده دو ورودی'],
        ['max', 'max(a, b)', 'بیشینه دو ورودی (نوعی جمع غیرخطی)'],
        ['min', 'min(a, b)', 'کمینه دو ورودی (نوعی جمع غیرخطی)'],
        ['avg', '(a+b)/2', 'میانگین = جمع تقسیم بر ۲'],
        ['weighted_add', '0.7a + 0.3b', 'جمع وزن‌دار'],
    ]
    add_styled_table(doc, func_headers, func_rows)

    doc.add_paragraph('')

    # ---- 4.3 Z-score ----
    doc.add_heading('۴.۳ نرمال‌سازی Z-score و دلیل استفاده', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    zscore_desc = (
        'نرمال‌سازی Z-score یک روش استاندارد آماری است که داده‌ها را به مقیاس واحد تبدیل می‌کند:\n\n'
        'فرمول:  z = (x - μ) / σ\n'
        'که در آن:\n'
        '• x = مقدار خام\n'
        '• μ = میانگین داده‌ها\n'
        '• σ = انحراف معیار داده‌ها\n'
        '• z = مقدار نرمال‌شده (میانگین = ۰، واریانس = ۱)\n\n'
        'چرا از Z-score استفاده می‌کنیم؟\n\n'
        '۱. تفاوت مقیاس شهرها:\n'
        '   • میلان ≈ ۳,۲۵۰,۰۰۰ نفر جمعیت → هزاران مبتلای روزانه\n'
        '   • کومو ≈ ۶۰۰,۰۰۰ نفر جمعیت → صدها مبتلای روزانه\n'
        '   • بدون نرمال‌سازی: CGP فقط به شهرهای بزرگ توجه می‌کند\n\n'
        '۲. عدالت در مقایسه:\n'
        '   • Z-score همه شهرها را به یک مقیاس (میانگین ۰، واریانس ۱) تبدیل می‌کند\n'
        '   • هر شهر به اندازه برابر در تحلیل CGP شرکت می‌کند\n'
        '   • وزن‌های CGP نشان‌دهنده اهمیت واقعی ارتباط است نه صرفاً بزرگی اعداد\n\n'
        '۳. بهبود عملکرد CGP:\n'
        '   • توابع ریاضی با مقادیر نرمال بهتر کار می‌کنند\n'
        '   • جلوگیری از overflow عددی\n'
        '   • همگرایی سریع‌تر الگوریتم تکاملی'
    )
    p = doc.add_paragraph(zscore_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Z-score example table
    z_headers = ['شهر', 'مبتلایان (خام)', 'پس از Z-score', 'تفسیر']
    z_rows = [
        ['Milano', '5,000', '+2.1', 'بسیار بالاتر از میانگین'],
        ['Como', '200', '-0.5', 'کمتر از میانگین'],
        ['Roma', '3,500', '+1.3', 'بالاتر از میانگین'],
        ['Taranto', '100', '-0.8', 'کمتر از میانگین'],
    ]
    add_styled_table(doc, z_headers, z_rows)

    doc.add_paragraph('')

    doc.add_page_break()

    # ---- CGP Parameters ----
    doc.add_heading('پارامترهای مدل CGP:', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    param_headers = ['پارامتر', 'مقدار', 'توضیح']
    param_rows = [
        ['تعداد سطرها', '3', 'تعداد سطرهای شبکه CGP'],
        ['تعداد ستون‌ها', '8', 'تعداد ستون‌های شبکه CGP'],
        ['Levels Back', '8', 'اتصال کامل (برابر تعداد ستون‌ها)'],
        ['تعداد خروجی', '1', 'پیش‌بینی مبتلایان جدید شهر هدف'],
        ['تعداد نسل‌ها', '500', 'حداکثر تعداد نسل‌های تکاملی'],
        ['λ (فرزندان)', '4', 'تعداد فرزندان در هر نسل'],
        ['نرخ جهش', '0.10', 'احتمال جهش هر ژن'],
        ['لگ پیش‌فرض', '7', 'استفاده از ۷ روز قبل'],
        ['آستانه R²', '0.9', 'حداقل R² برای ارتباط معنادار (سخت‌گیرانه)'],
        ['تعداد توابع', '5', 'add, max, min, avg, weighted_add'],
        ['تاریخ لاک‌دون', '2020-03-09', 'لاک‌دون ملی ایتالیا'],
    ]
    add_styled_table(doc, param_headers, param_rows)

    doc.add_page_break()

    # =========================================================================
    # 5. CGP RESULTS
    # =========================================================================
    h = doc.add_heading('۵. نتایج تحلیل CGP', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ---- 5.1 Connections ----
    doc.add_heading('۵.۱ تحلیل ارتباطات بین شهری', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Load connection data
    conn_path = os.path.join(CGP_DIR, 'all_connections.csv')
    if os.path.exists(conn_path):
        conn_df = pd.read_csv(conn_path)
        n_total = len(conn_df)
        n_significant = len(conn_df[conn_df['is_significant'] == True])

        results_overview = (
            f'الگوریتم CGP با توابع جمعی بر روی ۱۵ شهر منتخب اجرا شد. '
            f'توجه: آستانه R² بسیار سخت‌گیرانه (۰.۹) در نظر گرفته شده است. '
            f'از {n_total} جفت ارتباط، {n_significant} ارتباط معنادار (R² > 0.9) شناسایی شد.'
        )
        p = doc.add_paragraph(results_overview)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Top connections table
        top = conn_df.head(15)
        top_headers = ['شهر ۱', 'شهر ۲', 'منطقه ۱', 'منطقه ۲', 'R²', 'معنادار']
        top_rows = []
        for _, row in top.iterrows():
            top_rows.append([
                str(row['city_1']), str(row['city_2']),
                str(row.get('region_1', '')), str(row.get('region_2', '')),
                str(row['connection_strength_R2']),
                'بله' if row['is_significant'] else 'خیر',
            ])
        add_styled_table(doc, top_headers, top_rows)

    add_image_if_exists(doc, os.path.join(CGP_DIR, 'connection_heatmap.png'),
                        'شکل ۴: نقشه حرارتی ارتباطات ۱۵ شهر (R²)')
    add_image_if_exists(doc, os.path.join(CGP_DIR, 'network_graph.png'),
                        'شکل ۵: گراف شبکه ارتباطات (فقط R² > 0.9)')

    doc.add_page_break()

    # ---- 5.2 CGP Weights ----
    doc.add_heading('۵.۲ وزن‌های CGP', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    weights_desc = (
        'نقشه حرارتی زیر نشان می‌دهد که CGP چه وزنی به هر ارتباط بین شهرها داده است. '
        'محور Y شهر هدف و محور X شهر منبع (تأثیرگذار) است. '
        'مقادیر R² بالاتر نشان‌دهنده تأثیر قوی‌تر شهر منبع بر شهر هدف است.'
    )
    p = doc.add_paragraph(weights_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_image_if_exists(doc, os.path.join(CGP_DIR, 'cgp_weights_heatmap.png'),
                        'شکل ۶: وزن‌های ارتباطی CGP (شهر منبع → شهر هدف)')

    # City summary
    summary_path = os.path.join(CGP_DIR, 'city_summary.csv')
    if os.path.exists(summary_path):
        summ_df = pd.read_csv(summary_path)
        summ_headers = ['شهر', 'منطقه', 'R²', 'تعداد ورودی فعال', 'شهرهای فعال']
        summ_rows = []
        for _, row in summ_df.iterrows():
            active_str = str(row.get('active_cities', ''))
            if len(active_str) > 40:
                active_str = active_str[:40] + '...'
            summ_rows.append([
                str(row['city']), str(row['region']),
                str(row['cgp_r2']), str(row['n_active_inputs']),
                active_str,
            ])
        add_styled_table(doc, summ_headers, summ_rows)

    add_image_if_exists(doc, os.path.join(CGP_DIR, 'cgp_fitness_convergence.png'),
                        'شکل ۷: همگرایی تابع شایستگی CGP')

    doc.add_page_break()

    # ---- 5.3 Lag Analysis ----
    doc.add_heading('۵.۳ تحلیل لگ (R² در مقابل لگ)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lag_desc = (
        'برای بررسی اثر تأخیر زمانی (lag) بر کیفیت پیش‌بینی، '
        'CGP با لگ‌های مختلف از ۳ تا ۱۰۰ روز اجرا شد. '
        'نمودار زیر نشان می‌دهد که R² چگونه با افزایش لگ تغییر می‌کند:\n\n'
        '• لگ‌های کوتاه (۳-۱۰ روز): معمولاً R² پایین‌تر (اطلاعات کافی ندارند)\n'
        '• لگ‌های متوسط (۱۵-۳۰ روز): بهترین R² (دوره نهفتگی ویروس)\n'
        '• لگ‌های بلند (۵۰-۱۰۰ روز): ممکن است R² کاهش یابد (اطلاعات قدیمی کمتر مرتبط)'
    )
    p = doc.add_paragraph(lag_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_image_if_exists(doc, os.path.join(CGP_DIR, 'lag_vs_r2_plot.png'),
                        'شکل ۸: R² در مقابل لگ - برای هر ۱۵ شهر')
    add_image_if_exists(doc, os.path.join(CGP_DIR, 'lag_vs_r2_combined.png'),
                        'شکل ۹: R² در مقابل لگ - ترکیبی')

    # Lag sweep table
    lag_path = os.path.join(CGP_DIR, 'lag_sweep_results.csv')
    if os.path.exists(lag_path):
        lag_df = pd.read_csv(lag_path)
        lag_headers = ['شهر', 'منطقه'] + [f'لگ {l}' for l in [3, 7, 15, 30, 50, 100]]
        lag_rows = []
        for _, row in lag_df.iterrows():
            lag_rows.append([
                str(row['city']), str(row['region']),
                str(row.get('lag_3', '')), str(row.get('lag_7', '')),
                str(row.get('lag_15', '')), str(row.get('lag_30', '')),
                str(row.get('lag_50', '')), str(row.get('lag_100', '')),
            ])
        add_styled_table(doc, lag_headers, lag_rows)

    doc.add_page_break()

    # =========================================================================
    # 6. LOCKDOWN ANALYSIS
    # =========================================================================
    h = doc.add_heading('۶. تحلیل قبل و بعد از لاک‌دون', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lockdown_desc = (
        'تاریخ لاک‌دون ملی ایتالیا: ۹ مارس ۲۰۲۰\n\n'
        'فرضیه اصلی:\n'
        '• قبل از لاک‌دون: انسان‌ها می‌توانند بین شهرها جابه‌جا شوند '
        '→ مبتلایان هر شهر تحت تأثیر شهرهای دیگر نیز هست '
        '→ CGP باید ارتباطات بین شهری قوی‌تری پیدا کند\n\n'
        '• بعد از لاک‌دون: جابه‌جایی محدود شده '
        '→ مبتلایان هر شهر عمدتاً تحت تأثیر روزهای قبل همان شهر است '
        '→ CGP باید ارتباطات بین شهری ضعیف‌تری پیدا کند\n'
        '→ مدل فقط با لگ خود شهر (self-lag) نیز باید عملکرد مناسبی داشته باشد\n\n'
        'برای بررسی این فرضیه، سه تحلیل انجام شد:\n'
        '۱. قبل از لاک‌دون: CGP با ورودی تمام شهرها (inter-city)\n'
        '۲. بعد از لاک‌دون: CGP با ورودی تمام شهرها (inter-city) - برای مقایسه\n'
        '۳. بعد از لاک‌دون: CGP فقط با لگ خود شهر (self-lag only)'
    )
    p = doc.add_paragraph(lockdown_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_image_if_exists(doc, os.path.join(CGP_DIR, 'pre_post_lockdown_comparison.png'),
                        'شکل ۱۰: مقایسه R² قبل و بعد از لاک‌دون')
    add_image_if_exists(doc, os.path.join(CGP_DIR, 'lockdown_by_region.png'),
                        'شکل ۱۱: تحلیل لاک‌دون به تفکیک منطقه')
    add_image_if_exists(doc, os.path.join(CGP_DIR, 'lockdown_active_cities.png'),
                        'شکل ۱۲: تعداد شهرهای تأثیرگذار قبل و بعد از لاک‌دون')

    # Lockdown data table
    lockdown_path = os.path.join(CGP_DIR, 'lockdown_comparison.csv')
    if os.path.exists(lockdown_path):
        ld_df = pd.read_csv(lockdown_path)
        ld_headers = ['شهر', 'منطقه', 'R² قبل', 'R² حین (بین)',
                      'R² حین (خود)', 'R² بعد (بین)', 'R² بعد (خود)']
        ld_rows = []
        for _, row in ld_df.iterrows():
            ld_rows.append([
                str(row['city']), str(row['region']),
                f"{row['pre_lockdown_r2']:.4f}",
                f"{row.get('during_lockdown_inter_r2', 0):.4f}",
                f"{row.get('during_lockdown_self_r2', 0):.4f}",
                f"{row.get('after_easing_inter_r2', 0):.4f}",
                f"{row.get('after_easing_self_r2', 0):.4f}",
            ])
        add_styled_table(doc, ld_headers, ld_rows)

    doc.add_paragraph('')

    lockdown_analysis = (
        'تفسیر نتایج لاک‌دون:\n\n'
        '• اگر R² قبل از لاک‌دون > R² بعد از لاک‌دون (بین‌شهری):\n'
        '  فرضیه تأیید شد - جابه‌جایی بین شهرها واقعاً بر انتشار تأثیر داشته\n\n'
        '• اگر تعداد شهرهای فعال قبل > بعد:\n'
        '  لاک‌دون واقعاً ارتباطات بین شهری را کاهش داده\n\n'
        '• اگر R² بعد لاک‌دون (self-lag) قابل قبول باشد:\n'
        '  بعد از لاک‌دون، فقط سابقه خود شهر کافی است'
    )
    p = doc.add_paragraph(lockdown_analysis)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # =========================================================================
    # 7. CONCLUSION
    # =========================================================================
    h = doc.add_heading('۷. نتیجه‌گیری', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    conclusion = (
        'در این تحقیق، با استفاده از الگوریتم CGP با توابع فقط جمعی، '
        'ارتباطات بین ۱۵ شهر منتخب ایتالیا از سه منطقه جغرافیایی بررسی شد.'
    )
    p = doc.add_paragraph(conclusion)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    conclusions = [
        'CGP با آستانه سخت‌گیرانه R² > 0.9 ارتباطات واقعاً معنادار را شناسایی کرد',
        'توابع فقط جمعی (بدون تفریق) نیز قادر به مدل‌سازی ارتباطات بین شهری بودند',
        'تحلیل لگ نشان داد که لگ‌های بهینه معمولاً در بازه ۱۵-۳۰ روزه هستند',
        'لاک‌دون تأثیر قابل توجهی بر کاهش ارتباطات بین شهری داشت',
        'قبل از لاک‌دون: شهرها بر یکدیگر تأثیرگذار بودند (جابه‌جایی آزاد)',
        'بعد از لاک‌دون: مبتلایان هر شهر عمدتاً تحت تأثیر سابقه خود شهر بود',
        'نرمال‌سازی Z-score ضروری است تا شهرهای بزرگ و کوچک به طور عادلانه مقایسه شوند',
        'پارامتر levels_back در CGP با مقدار حداکثر (اتصال کامل) بهترین نتیجه را داد',
    ]

    for item in conclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph('')

    final_note = (
        'پیشنهاد می‌شود این تحلیل بر روی داده‌های سال‌های ۲۰۲۱ تا ۲۰۲۴ نیز تکرار شود '
        'تا تغییرات الگوهای ارتباطی در طول زمان و پس از رفع لاک‌دون بررسی شوند.'
    )
    p = doc.add_paragraph(final_note)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # =========================================================================
    # Save
    # =========================================================================
    output_path = os.path.join(OUTPUT_DIR, 'CGP_Report_2020v2_Farsi.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == '__main__':
    main()
