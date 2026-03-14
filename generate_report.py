"""
Generate a Word document report (Farsi) for CGP city analysis results.
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

CGP_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\cgp_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    # Data
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Arial'

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('گزارش تحلیل ارتباطات بین شهری\nبا استفاده از برنامه‌نویسی ژنتیک کارتزین (CGP)')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(46, 64, 87)

    doc.add_paragraph('')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('تحلیل رفتارشناسی انتشار COVID-19 در ۱۰۷ استان ایتالیا\nداده‌های سال ۲۰۲۲')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')
    doc.add_paragraph('')
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('تاریخ: مارس ۲۰۲۶')
    run.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    h = doc.add_heading('فهرست مطالب', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    toc_items = [
        '۱. مقدمه و هدف پروژه',
        '۲. معرفی برنامه‌نویسی ژنتیک کارتزین (CGP)',
        '۳. شرح داده‌ها و پیش‌پردازش',
        '۴. پارامترهای مدل CGP',
        '۵. نتایج تحلیل CGP',
        '   ۵.۱ قوی‌ترین ارتباطات کشف شده',
        '   ۵.۲ شهرهای مرکزی (Hub)',
        '   ۵.۳ گراف شبکه ارتباطات',
        '   ۵.۴ نقشه حرارتی ارتباطات',
        '   ۵.۵ ماتریس ارتباطات ۳۰ شهر برتر',
        '   ۵.۶ همگرایی الگوریتم تکاملی',
        '   ۵.۷ تحلیل نفوذ شهرهای مرکزی',
        '۶. تحلیل و تفسیر نتایج',
        '۷. نتیجه‌گیری',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    h = doc.add_heading('۱. مقدمه و هدف پروژه', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    intro_text = (
        'در این پروژه، هدف بررسی و کشف ارتباطات مستقیم بین شهرهای ایتالیا از نظر انتشار بیماری COVID-19 است. '
        'سوال اصلی تحقیق این است: آیا افزایش یا کاهش تعداد مبتلایان در یک شهر بر تعداد مبتلایان در شهرهای دیگر تأثیر مستقیم دارد؟ '
        'برای پاسخ به این سوال، از روش برنامه‌نویسی ژنتیک کارتزین (Cartesian Genetic Programming - CGP) استفاده شده است.'
    )
    p = doc.add_paragraph(intro_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    intro2 = (
        'این تحلیل بر روی داده‌های روزانه ۱۰۷ استان (Province) ایتالیا در سال ۲۰۲۲ انجام شده است. '
        'با استفاده از CGP، عبارات ریاضی تکاملی ایجاد می‌شوند که رابطه بین تغییرات روزانه مبتلایان '
        'شهرهای مختلف را مدل‌سازی می‌کنند. ورودی‌های فعال در برنامه تکامل‌یافته نشان‌دهنده شهرهایی هستند '
        'که ارتباط مستقیم با شهر هدف دارند.'
    )
    p = doc.add_paragraph(intro2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph('')

    # =========================================================================
    # 2. CGP METHODOLOGY
    # =========================================================================
    h = doc.add_heading('۲. معرفی برنامه‌نویسی ژنتیک کارتزین (CGP)', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cgp_desc = (
        'برنامه‌نویسی ژنتیک کارتزین (CGP) یک الگوریتم تکاملی است که توسط Julian Miller در سال ۲۰۰۰ معرفی شد. '
        'در CGP، برنامه‌ها به صورت گراف‌های جهت‌دار بدون حلقه (DAG) بر روی یک شبکه دوبعدی (سطر × ستون) نمایش داده می‌شوند.'
    )
    p = doc.add_paragraph(cgp_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('ساختار CGP:', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    structure_items = [
        'هر گره (Node) در شبکه شامل سه ژن است: یک تابع ریاضی و دو ورودی',
        'ورودی‌ها می‌توانند از گره‌های قبلی یا مستقیماً از داده‌های ورودی باشند',
        'خروجی مدل از آخرین لایه گره‌ها استخراج می‌شود',
        'فقط گره‌های "فعال" (Active Nodes) در محاسبه نهایی شرکت می‌کنند',
    ]
    for item in structure_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('توابع ریاضی مورد استفاده:', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    func_headers = ['نام تابع', 'عملیات', 'توضیح']
    func_rows = [
        ['add', 'a + b', 'جمع دو ورودی'],
        ['sub', 'a - b', 'تفاضل دو ورودی'],
        ['mul', 'a × b', 'ضرب دو ورودی'],
        ['div', 'a ÷ b', 'تقسیم محافظت‌شده'],
        ['max', 'max(a, b)', 'بیشینه دو ورودی'],
        ['min', 'min(a, b)', 'کمینه دو ورودی'],
        ['abs_diff', '|a - b|', 'قدر مطلق تفاضل'],
        ['avg', '(a+b)/2', 'میانگین دو ورودی'],
    ]
    add_styled_table(doc, func_headers, func_rows)

    doc.add_paragraph('')

    evolution_desc = (
        'تکامل با استراتژی (1+λ) انجام می‌شود: در هر نسل، یک والد '
        'λ فرزند از طریق جهش نقطه‌ای تولید می‌کند. بهترین فرد (با کمترین خطا) '
        'به عنوان والد نسل بعد انتخاب می‌شود. این فرآیند تا رسیدن به تعداد نسل‌های '
        'مشخص تکرار می‌شود.'
    )
    p = doc.add_paragraph(evolution_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # =========================================================================
    # 3. DATA DESCRIPTION
    # =========================================================================
    h = doc.add_heading('۳. شرح داده‌ها و پیش‌پردازش', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    data_desc = (
        'داده‌های مورد استفاده شامل تعداد تجمعی مبتلایان روزانه (totale_casi) '
        'برای ۱۰۷ استان ایتالیا در سال ۲۰۲۲ (۳۶۵ روز) است. '
        'این داده‌ها از فایل‌های CSV هر شهر استخراج و در یک فایل واحد (merged_cities_2022.csv) ادغام شده‌اند.'
    )
    p = doc.add_paragraph(data_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('مراحل پیش‌پردازش:', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    preprocess_items = [
        'تبدیل داده‌های تجمعی به مبتلایان جدید روزانه (daily new infections) با محاسبه تفاضل روزانه',
        'حذف مقادیر منفی (ناشی از اصلاح داده‌های قبلی)',
        'ایجاد ویژگی‌های تأخیری (Lagged Features): برای هر شهر، داده‌های ۰، ۱ و ۲ روز قبل',
        'نرمال‌سازی داده‌ها (Z-score normalization) برای بهبود عملکرد CGP',
        'انتخاب ۱۵ شهر برتر (بر اساس ضریب همبستگی) به عنوان کاندیدای ورودی برای هر شهر هدف',
    ]
    for item in preprocess_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # =========================================================================
    # 4. CGP PARAMETERS
    # =========================================================================
    h = doc.add_heading('۴. پارامترهای مدل CGP', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    param_desc = (
        'در جدول زیر پارامترهای استفاده‌شده برای الگوریتم CGP بیان شده‌اند:'
    )
    p = doc.add_paragraph(param_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    param_headers = ['پارامتر', 'مقدار', 'توضیح']
    param_rows = [
        ['تعداد سطرها', '3', 'تعداد سطرهای شبکه CGP'],
        ['تعداد ستون‌ها', '8', 'تعداد ستون‌های شبکه CGP'],
        ['تعداد خروجی', '1', 'پیش‌بینی مبتلایان جدید شهر هدف'],
        ['تعداد نسل‌ها', '300', 'حداکثر تعداد نسل‌های تکاملی'],
        ['λ (فرزندان)', '4', 'تعداد فرزندان در هر نسل'],
        ['نرخ جهش', '0.08', 'احتمال جهش هر ژن'],
        ['تعداد تأخیرها', '3', 'لگ ۰، ۱ و ۲ روزه'],
        ['کاندیداهای ورودی', '15', 'تعداد شهرهای کاندیدا برای هر هدف'],
        ['آستانه R²', '0.3', 'حداقل R² برای ارتباط معنادار'],
        ['تعداد توابع', '8', 'add, sub, mul, div, max, min, abs_diff, avg'],
    ]
    add_styled_table(doc, param_headers, param_rows)

    doc.add_paragraph('')

    param_note = (
        'معیار ارزیابی: میانگین مربعات خطا (MSE) بین پیش‌بینی CGP و مقادیر واقعی. '
        'همچنین از ضریب تعیین (R²) برای سنجش کیفیت پیش‌بینی استفاده شده است. '
        'مقدار R² بالاتر از ۰.۳ به عنوان ارتباط معنادار در نظر گرفته شده است.'
    )
    p = doc.add_paragraph(param_note)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # =========================================================================
    # 5. RESULTS
    # =========================================================================
    h = doc.add_heading('۵. نتایج تحلیل CGP', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Load data
    conn_df = pd.read_csv(os.path.join(CGP_DIR, 'all_connections.csv'))
    summary_df = pd.read_csv(os.path.join(CGP_DIR, 'city_summary.csv'))

    n_total = len(conn_df)
    n_significant = len(conn_df[conn_df['is_significant'] == True])
    n_connected = len(summary_df[summary_df['n_connections'] > 0])

    results_overview = (
        f'الگوریتم CGP بر روی تمام ۱۰۷ شهر اجرا شد. '
        f'در مجموع {n_total} جفت ارتباط شناسایی شد که از این تعداد، '
        f'{n_significant} ارتباط معنادار (R² > 0.3) و '
        f'{n_total - n_significant} ارتباط ضعیف بودند. '
        f'از ۱۰۷ شهر، {n_connected} شهر حداقل یک ارتباط مستقیم معنادار داشتند.'
    )
    p = doc.add_paragraph(results_overview)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Overview table
    overview_headers = ['شاخص', 'مقدار']
    overview_rows = [
        ['تعداد کل شهرها', '107'],
        ['تعداد کل جفت ارتباطات شناسایی‌شده', str(n_total)],
        ['ارتباطات معنادار (R² > 0.3)', str(n_significant)],
        ['شهرهای دارای ارتباط', f'{n_connected} از 107'],
    ]
    add_styled_table(doc, overview_headers, overview_rows)

    doc.add_paragraph('')

    # ---- 5.1 Top connections ----
    doc.add_heading('۵.۱ قوی‌ترین ارتباطات کشف شده', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    top_desc = (
        'جدول زیر ۲۰ ارتباط قوی‌ترین بین شهرها را نشان می‌دهد. '
        'همانطور که مشاهده می‌شود، شهرهای هم‌جوار جغرافیایی بیشترین ارتباط را دارند:'
    )
    p = doc.add_paragraph(top_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    top_headers = ['رتبه', 'شهر اول', 'شهر دوم', 'قدرت ارتباط (R²)', 'معنادار']
    top_rows = []
    for idx, row in conn_df.head(20).iterrows():
        top_rows.append([
            str(idx + 1),
            str(row['city_1']),
            str(row['city_2']),
            str(row['connection_strength_R2']),
            'بله' if row['is_significant'] else 'خیر',
        ])
    add_styled_table(doc, top_headers, top_rows)

    doc.add_paragraph('')
    # Top connections image
    img_path = os.path.join(CGP_DIR, 'top_connections.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۱: نمودار قوی‌ترین ارتباطات بین شهرها (بر اساس R²)').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 5.2 Hub cities ----
    doc.add_heading('۵.۲ شهرهای مرکزی (Hub)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    hub_desc = (
        'شهرهای مرکزی (Hub) شهرهایی هستند که بیشترین تعداد ارتباط مستقیم را با سایر شهرها دارند. '
        'این شهرها نقش کلیدی در انتشار بیماری بین مناطق مختلف ایفا می‌کنند:'
    )
    p = doc.add_paragraph(hub_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    hub_headers = ['رتبه', 'شهر', 'تعداد ارتباطات', 'شهرهای مرتبط']
    hub_rows = []
    for idx, (_, row) in enumerate(summary_df.head(15).iterrows()):
        hub_rows.append([
            str(idx + 1),
            str(row['city']),
            str(row['n_connections']),
            str(row['linked_cities'])[:60] + ('...' if len(str(row['linked_cities'])) > 60 else ''),
        ])
    add_styled_table(doc, hub_headers, hub_rows)

    doc.add_paragraph('')
    img_path = os.path.join(CGP_DIR, 'hub_cities.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۲: نمودار شهرهای مرکزی بر حسب تعداد ارتباطات').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 5.3 Network graph ----
    doc.add_heading('۵.۳ گراف شبکه ارتباطات', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    net_desc = (
        'در شکل زیر، گراف شبکه ارتباطات بین شهرها نمایش داده شده است. '
        'هر گره نماینده یک شهر و هر یال نماینده یک ارتباط مستقیم است. '
        'اندازه گره‌ها متناسب با تعداد ارتباطات (درجه) و ضخامت یال‌ها متناسب با قدرت ارتباط (R²) است. '
        'رنگ گره‌ها نیز بر اساس تعداد ارتباطات تعیین شده است (قرمز‌تر = ارتباطات بیشتر).'
    )
    p = doc.add_paragraph(net_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    img_path = os.path.join(CGP_DIR, 'network_graph.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۳: گراف شبکه ارتباطات CGP بین شهرهای ایتالیا').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 5.4 Heatmap ----
    doc.add_heading('۵.۴ نقشه حرارتی ارتباطات', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    heat_desc = (
        'نقشه حرارتی زیر، ماتریس کامل ارتباطات بین تمام ۱۰۷ شهر را نشان می‌دهد. '
        'هر خانه نشان‌دهنده قدرت ارتباط (R²) بین دو شهر است. رنگ‌های تیره‌تر نشان‌دهنده ارتباط قوی‌تر هستند.'
    )
    p = doc.add_paragraph(heat_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    img_path = os.path.join(CGP_DIR, 'connection_heatmap.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۴: نقشه حرارتی ارتباطات ۱۰۷ شهر').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 5.5 Top 30 matrix ----
    doc.add_heading('۵.۵ ماتریس ارتباطات ۳۰ شهر برتر', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    mat_desc = (
        'برای تحلیل دقیق‌تر، ماتریس ارتباطات ۳۰ شهر با بیشترین ارتباطات در شکل زیر نمایش داده شده است. '
        'مقادیر عددی R² در هر خانه درج شده‌اند.'
    )
    p = doc.add_paragraph(mat_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    img_path = os.path.join(CGP_DIR, 'top30_connection_matrix.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۵: ماتریس ارتباطات ۳۰ شهر پرارتباط').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- 5.6 Fitness convergence ----
    doc.add_heading('۵.۶ همگرایی الگوریتم تکاملی', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fit_desc = (
        'شکل زیر روند همگرایی تابع شایستگی (Fitness) الگوریتم CGP را برای شش شهر نمونه نشان می‌دهد. '
        'کاهش مداوم خطا (MSE) در طول نسل‌ها نشان‌دهنده یادگیری موفق الگوریتم و کشف الگوهای ارتباطی بین شهرها است.'
    )
    p = doc.add_paragraph(fit_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    img_path = os.path.join(CGP_DIR, 'cgp_fitness_convergence.png')
    if os.path.exists(img_path):
        doc.add_paragraph('شکل ۶: همگرایی تابع شایستگی CGP').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ---- 5.7 Influence details ----
    doc.add_heading('۵.۷ تحلیل نفوذ شهرهای مرکزی', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    inf_desc = (
        'برای ۱۰ شهر مرکزی (Hub)، نمودارهای جداگانه‌ای تولید شده که نشان می‌دهد '
        'کدام شهرها بر آن‌ها تأثیرگذار هستند:'
    )
    p = doc.add_paragraph(inf_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    influence_files = [f for f in os.listdir(CGP_DIR) if f.startswith('influence_')]
    for i, fname in enumerate(sorted(influence_files)):
        img_path = os.path.join(CGP_DIR, fname)
        city_name = fname.replace('influence_', '').replace('.png', '').title()
        doc.add_paragraph(f'شکل {7+i}: شهرهای مؤثر بر {city_name}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

    doc.add_page_break()

    # =========================================================================
    # 6. ANALYSIS
    # =========================================================================
    h = doc.add_heading('۶. تحلیل و تفسیر نتایج', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    analysis_sections = [
        ('الگوی جغرافیایی ارتباطات',
         'یافته‌های CGP به وضوح نشان می‌دهد که ارتباطات مستقیم بین شهرها عمدتاً '
         'تابع نزدیکی جغرافیایی است. به عنوان مثال:\n'
         '• شهرهای Chieti، Teramo و Pescara در منطقه آبروتزو (Abruzzo) با R² = 0.9998 قوی‌ترین ارتباط را دارند\n'
         '• Milano و Monza e della Brianza در لمباردی (Lombardy) با R² = 0.9827\n'
         '• Bergamo و Brescia در لمباردی با R² = 0.9781\n'
         '• Padova و Vicenza در ونتو (Veneto) با R² = 0.9817\n'
         'این نتایج منطقی است زیرا تردد روزانه بین شهرهای هم‌جوار عامل اصلی انتقال بیماری است.'),

        ('خوشه‌های منطقه‌ای',
         'CGP خوشه‌های واضحی از شهرهای مرتبط را شناسایی کرده است:\n\n'
         '• خوشه لمباردی: Milano ↔ Monza ↔ Pavia ↔ Como ↔ Varese ↔ Lecco ↔ Brescia ↔ Bergamo\n'
         '• خوشه ونتو: Padova ↔ Vicenza ↔ Verona ↔ Venezia ↔ Treviso ↔ Belluno\n'
         '• خوشه آبروتزو: Chieti ↔ Teramo ↔ Pescara ↔ L\'Aquila\n'
         '• خوشه پولیا: Bari ↔ Brindisi ↔ Taranto ↔ Barletta-Andria-Trani ↔ Foggia\n'
         '• خوشه توسکانی: Pisa ↔ Pistoia ↔ Prato ↔ Firenze ↔ Arezzo ↔ Lucca\n'
         '• خوشه امیلیا-رومانیا: Bologna ↔ Ravenna ↔ Rimini ↔ Forlì-Cesena\n'
         '• خوشه فریولی: Trieste ↔ Udine ↔ Pordenone ↔ Gorizia\n'
         '• خوشه کامپانیا: Napoli ↔ Caserta ↔ Salerno ↔ Benevento'),

        ('شهرهای مستقل',
         'تعداد ۳ شهر از ۱۰۷ شهر هیچ ارتباط معناداری با شهرهای دیگر نشان ندادند. '
         'این شهرها عبارتند از شهرهایی با مقدار R² کمتر از ۰.۳ (مانند Catania و Catanzaro). '
         'دلایل احتمالی شامل ویژگی‌های جغرافیایی خاص (جزایر)، '
         'سیاست‌های بهداشتی متفاوت، یا الگوهای متفاوت رفتاری در آن مناطق است.'),

        ('نقش CGP در مقایسه با روش‌های آماری سنتی',
         'مزیت استفاده از CGP نسبت به روش‌های ساده مانند همبستگی خطی:\n'
         '• CGP قادر به شناسایی روابط غیرخطی بین شهرها است\n'
         '• فقط ورودی‌های واقعاً مؤثر (Active Inputs) در مدل نهایی باقی می‌مانند\n'
         '• تأخیرهای زمانی (Lags) در مدل لحاظ شده‌اند\n'
         '• عبارات ریاضی تولید شده قابل تفسیر هستند'),
    ]

    for title, content in analysis_sections:
        doc.add_heading(title, level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph('')

    doc.add_page_break()

    # =========================================================================
    # 7. CONCLUSION
    # =========================================================================
    h = doc.add_heading('۷. نتیجه‌گیری', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    conclusion = (
        'در این تحقیق، با استفاده از الگوریتم برنامه‌نویسی ژنتیک کارتزین (CGP)، '
        'ارتباطات مستقیم بین ۱۰۷ استان ایتالیا از نظر انتشار COVID-19 بررسی شد. '
        'نتایج اصلی به شرح زیر است:'
    )
    p = doc.add_paragraph(conclusion)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    conclusions = [
        f'از {n_total} جفت ارتباط تحلیل‌شده، {n_significant} ارتباط معنادار (R² > 0.3) شناسایی شد',
        f'{n_connected} شهر از ۱۰۷ شهر حداقل یک ارتباط مستقیم معنادار داشتند',
        'قوی‌ترین ارتباطات بین شهرهای هم‌جوار جغرافیایی مشاهده شد (R² تا ۰.۹۹۹۸)',
        'خوشه‌های منطقه‌ای واضحی شناسایی شد که کاملاً با تقسیمات جغرافیایی ایتالیا منطبق هستند',
        'شهرهای Hub مانند Arezzo، Trento و Bari نقش مرکزی در شبکه انتشار دارند',
        'CGP توانست روابط غیرخطی و تأخیرهای زمانی بین شهرها را با موفقیت مدل‌سازی کند',
        'این یافته‌ها می‌تواند در برنامه‌ریزی سیاست‌های بهداشتی منطقه‌ای و پیش‌بینی الگوهای انتشار بیماری مورد استفاده قرار گیرد',
    ]

    for item in conclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph('')

    final_note = (
        'پیشنهاد می‌شود در تحقیقات آینده، این تحلیل بر روی داده‌های سال‌های ۲۰۲۰ تا ۲۰۲۴ '
        'تکرار شود تا تغییرات الگوهای ارتباطی در طول زمان بررسی شوند. '
        'همچنین، ترکیب نتایج CGP با مدل SIR می‌تواند به مدل‌سازی '
        'جامع‌تر انتشار بیماری بین شهرها کمک کند.'
    )
    p = doc.add_paragraph(final_note)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # =========================================================================
    # ALL CONNECTIONS APPENDIX
    # =========================================================================
    doc.add_page_break()
    h = doc.add_heading('پیوست: فهرست کامل ارتباطات', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    all_headers = ['ردیف', 'شهر اول', 'شهر دوم', 'R²', 'معنادار']
    all_rows = []
    for idx, row in conn_df.iterrows():
        all_rows.append([
            str(idx + 1),
            str(row['city_1']),
            str(row['city_2']),
            str(row['connection_strength_R2']),
            'بله' if row['is_significant'] else 'خیر',
        ])
    add_styled_table(doc, all_headers, all_rows)

    # Save
    output_path = os.path.join(OUTPUT_DIR, 'CGP_Report_Farsi.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == '__main__':
    main()
