"""
Generate an English Word document report for CGP city analysis results.
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

CGP_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\cgp_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Arial'

    # === TITLE PAGE ===
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Inter-City Infection Relationship Analysis\nUsing Cartesian Genetic Programming (CGP)')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(46, 64, 87)

    doc.add_paragraph('')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Behavioral Analysis of COVID-19 Spread Across 107 Italian Provinces\nYear 2022 Data')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')
    doc.add_paragraph('')
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('Date: March 2026')
    run.font.size = Pt(12)

    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction and Objectives',
        '2. Cartesian Genetic Programming (CGP) Overview',
        '3. Data Description and Preprocessing',
        '4. CGP Model Parameters',
        '5. CGP Analysis Results',
        '   5.1 Strongest Discovered Connections',
        '   5.2 Hub Cities',
        '   5.3 Network Graph',
        '   5.4 Connection Heatmap',
        '   5.5 Top 30 Cities Connection Matrix',
        '   5.6 Evolutionary Fitness Convergence',
        '   5.7 Hub City Influence Analysis',
        '6. Discussion and Interpretation',
        '7. Conclusions',
        'Appendix: Complete List of Connections',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # === 1. INTRODUCTION ===
    doc.add_heading('1. Introduction and Objectives', level=1)
    doc.add_paragraph(
        'This study investigates direct relationships between Italian cities regarding the spread '
        'of COVID-19 infections. The central research question is: does an increase or decrease '
        'in the number of infected individuals in one city directly influence infection counts in '
        'other cities? To answer this question, we employ Cartesian Genetic Programming (CGP), '
        'an evolutionary algorithm capable of discovering complex, non-linear mathematical '
        'relationships between time series data.'
    )
    doc.add_paragraph(
        'The analysis was conducted on daily infection data from 107 Italian provinces for the '
        'year 2022. CGP evolves mathematical expressions that model the relationship between '
        'daily new infection changes across different cities. The active inputs in the evolved '
        'programs reveal which cities have a direct influence on the target city\'s infection dynamics.'
    )

    # === 2. CGP METHODOLOGY ===
    doc.add_heading('2. Cartesian Genetic Programming (CGP) Overview', level=1)
    doc.add_paragraph(
        'Cartesian Genetic Programming (CGP) is an evolutionary algorithm introduced by Julian '
        'Miller in 2000. In CGP, programs are represented as Directed Acyclic Graphs (DAGs) '
        'arranged on a two-dimensional grid (rows x columns).'
    )

    doc.add_heading('CGP Structure:', level=3)
    items = [
        'Each node in the grid contains three genes: a mathematical function and two input connections',
        'Inputs can come from previous nodes or directly from the input data',
        'The model output is extracted from the last layer of nodes',
        'Only "Active Nodes" participate in the final computation (neutral code is ignored)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Mathematical Functions Used:', level=3)
    func_headers = ['Function', 'Operation', 'Description']
    func_rows = [
        ['add', 'a + b', 'Sum of two inputs'],
        ['sub', 'a - b', 'Difference of two inputs'],
        ['mul', 'a x b', 'Product of two inputs'],
        ['div', 'a / b', 'Protected division (avoids division by zero)'],
        ['max', 'max(a, b)', 'Maximum of two inputs'],
        ['min', 'min(a, b)', 'Minimum of two inputs'],
        ['abs_diff', '|a - b|', 'Absolute difference'],
        ['avg', '(a+b)/2', 'Average of two inputs'],
    ]
    add_styled_table(doc, func_headers, func_rows)

    doc.add_paragraph('')
    doc.add_paragraph(
        'Evolution follows a (1+lambda) strategy: in each generation, one parent produces lambda '
        'offspring through point mutation. The best individual (lowest error) is selected as the '
        'parent for the next generation. This process repeats for a specified number of generations.'
    )
    doc.add_page_break()

    # === 3. DATA DESCRIPTION ===
    doc.add_heading('3. Data Description and Preprocessing', level=1)
    doc.add_paragraph(
        'The dataset consists of daily cumulative infection counts (totale_casi) for 107 Italian '
        'provinces throughout 2022 (365 days). The data was extracted from individual city CSV files '
        'and merged into a single file (merged_cities_2022.csv).'
    )

    doc.add_heading('Preprocessing Steps:', level=3)
    preprocess = [
        'Converting cumulative data to daily new infections by computing first differences',
        'Removing negative values (artifacts from data corrections)',
        'Creating lagged features (lags 0, 1, and 2 days) for each candidate city',
        'Z-score normalization for improved CGP performance',
        'Selecting top 15 most correlated cities as CGP input candidates per target city',
    ]
    for item in preprocess:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # === 4. CGP PARAMETERS ===
    doc.add_heading('4. CGP Model Parameters', level=1)
    doc.add_paragraph('The following table summarizes the CGP parameters used in this analysis:')

    param_headers = ['Parameter', 'Value', 'Description']
    param_rows = [
        ['Grid Rows', '3', 'Number of rows in the CGP grid'],
        ['Grid Columns', '8', 'Number of columns in the CGP grid'],
        ['Outputs', '1', 'Prediction of target city daily new infections'],
        ['Generations', '300', 'Maximum number of evolutionary generations'],
        ['Lambda', '4', 'Number of offspring per generation'],
        ['Mutation Rate', '0.08', 'Probability of mutating each gene'],
        ['Number of Lags', '3', 'Lags 0, 1, and 2 days'],
        ['Input Candidates', '15', 'Top correlated cities per target'],
        ['R-squared Threshold', '0.3', 'Minimum R-squared for significant connection'],
        ['Number of Functions', '8', 'add, sub, mul, div, max, min, abs_diff, avg'],
    ]
    add_styled_table(doc, param_headers, param_rows)

    doc.add_paragraph('')
    doc.add_paragraph(
        'Evaluation metric: Mean Squared Error (MSE) between CGP predictions and actual values. '
        'The coefficient of determination (R-squared) is used to assess prediction quality. '
        'An R-squared value above 0.3 is considered a significant connection.'
    )
    doc.add_page_break()

    # === 5. RESULTS ===
    doc.add_heading('5. CGP Analysis Results', level=1)

    conn_df = pd.read_csv(os.path.join(CGP_DIR, 'all_connections.csv'))
    summary_df = pd.read_csv(os.path.join(CGP_DIR, 'city_summary.csv'))
    n_total = len(conn_df)
    n_significant = len(conn_df[conn_df['is_significant'] == True])
    n_connected = len(summary_df[summary_df['n_connections'] > 0])

    doc.add_paragraph(
        f'CGP was executed for all 107 cities. A total of {n_total} connection pairs were identified, '
        f'of which {n_significant} were significant (R-squared > 0.3) and '
        f'{n_total - n_significant} were weak. Out of 107 cities, {n_connected} had at least one '
        f'significant direct connection.'
    )

    overview_headers = ['Metric', 'Value']
    overview_rows = [
        ['Total Cities Analyzed', '107'],
        ['Total Connection Pairs Identified', str(n_total)],
        ['Significant Connections (R-squared > 0.3)', str(n_significant)],
        ['Cities with Connections', f'{n_connected} out of 107'],
    ]
    add_styled_table(doc, overview_headers, overview_rows)
    doc.add_paragraph('')

    # 5.1
    doc.add_heading('5.1 Strongest Discovered Connections', level=2)
    doc.add_paragraph(
        'The following table shows the top 20 strongest connections between cities. '
        'As observed, geographically adjacent cities exhibit the strongest relationships:'
    )

    top_headers = ['Rank', 'City 1', 'City 2', 'R-squared', 'Significant']
    top_rows = []
    for idx, row in conn_df.head(20).iterrows():
        top_rows.append([
            str(idx + 1), str(row['city_1']), str(row['city_2']),
            str(row['connection_strength_R2']),
            'Yes' if row['is_significant'] else 'No',
        ])
    add_styled_table(doc, top_headers, top_rows)
    doc.add_paragraph('')

    img = os.path.join(CGP_DIR, 'top_connections.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 1: Strongest inter-city connections (by R-squared)').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 5.2
    doc.add_heading('5.2 Hub Cities', level=2)
    doc.add_paragraph(
        'Hub cities are those with the highest number of direct connections to other cities. '
        'These cities play a key role in disease transmission between different regions:'
    )

    hub_headers = ['Rank', 'City', 'Connections', 'Linked Cities']
    hub_rows = []
    for idx, (_, row) in enumerate(summary_df.head(15).iterrows()):
        hub_rows.append([
            str(idx + 1), str(row['city']), str(row['n_connections']),
            str(row['linked_cities'])[:60] + ('...' if len(str(row['linked_cities'])) > 60 else ''),
        ])
    add_styled_table(doc, hub_headers, hub_rows)
    doc.add_paragraph('')

    img = os.path.join(CGP_DIR, 'hub_cities.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 2: Hub cities ranked by number of connections').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 5.3
    doc.add_heading('5.3 Network Graph', level=2)
    doc.add_paragraph(
        'The figure below shows the network graph of inter-city connections. Each node represents '
        'a city and each edge represents a direct connection. Node sizes are proportional to the '
        'number of connections (degree), edge thickness is proportional to connection strength (R-squared), '
        'and node colors indicate connection count (redder = more connections).'
    )

    img = os.path.join(CGP_DIR, 'network_graph.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 3: CGP-discovered city infection network').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 5.4
    doc.add_heading('5.4 Connection Heatmap', level=2)
    doc.add_paragraph(
        'The heatmap below shows the complete connection matrix for all 107 cities. '
        'Each cell represents the connection strength (R-squared) between two cities. '
        'Darker colors indicate stronger connections.'
    )

    img = os.path.join(CGP_DIR, 'connection_heatmap.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 4: Connection heatmap for 107 cities').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 5.5
    doc.add_heading('5.5 Top 30 Cities Connection Matrix', level=2)
    doc.add_paragraph(
        'For a more detailed analysis, the connection matrix for the 30 most connected cities '
        'is shown below with R-squared values annotated in each cell.'
    )

    img = os.path.join(CGP_DIR, 'top30_connection_matrix.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 5: Top 30 most connected cities matrix').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 5.6
    doc.add_heading('5.6 Evolutionary Fitness Convergence', level=2)
    doc.add_paragraph(
        'The figure below shows the fitness convergence (MSE reduction over generations) '
        'for six sample cities. The continuous decrease in error demonstrates successful '
        'learning and pattern discovery by the CGP algorithm.'
    )

    img = os.path.join(CGP_DIR, 'cgp_fitness_convergence.png')
    if os.path.exists(img):
        doc.add_paragraph('Figure 6: CGP fitness convergence').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # 5.7
    doc.add_heading('5.7 Hub City Influence Analysis', level=2)
    doc.add_paragraph(
        'For the top 10 hub cities, individual charts show which cities influence them:'
    )

    influence_files = sorted([f for f in os.listdir(CGP_DIR) if f.startswith('influence_')])
    for i, fname in enumerate(influence_files):
        img_path = os.path.join(CGP_DIR, fname)
        city_name = fname.replace('influence_', '').replace('.png', '').title()
        doc.add_paragraph(f'Figure {7+i}: Cities influencing {city_name}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
    doc.add_page_break()

    # === 6. DISCUSSION ===
    doc.add_heading('6. Discussion and Interpretation', level=1)

    sections = [
        ('Geographic Pattern of Connections',
         'CGP findings clearly demonstrate that direct connections between cities are predominantly '
         'a function of geographic proximity. Key examples include:\n\n'
         '  - Chieti, Teramo, and Pescara in Abruzzo: R-squared = 0.9998 (strongest)\n'
         '  - Milano and Monza e della Brianza in Lombardy: R-squared = 0.9827\n'
         '  - Bergamo and Brescia in Lombardy: R-squared = 0.9781\n'
         '  - Padova and Vicenza in Veneto: R-squared = 0.9817\n\n'
         'These results are logical since daily commuting between adjacent cities is the primary '
         'driver of disease transmission.'),

        ('Regional Clusters',
         'CGP identified clear clusters of related cities:\n\n'
         '  - Lombardy: Milano - Monza - Pavia - Como - Varese - Lecco - Brescia - Bergamo\n'
         '  - Veneto: Padova - Vicenza - Verona - Venezia - Treviso - Belluno\n'
         '  - Abruzzo: Chieti - Teramo - Pescara - L\'Aquila\n'
         '  - Puglia: Bari - Brindisi - Taranto - Barletta-Andria-Trani - Foggia\n'
         '  - Tuscany: Pisa - Pistoia - Prato - Firenze - Arezzo - Lucca\n'
         '  - Emilia-Romagna: Bologna - Ravenna - Rimini - Forli-Cesena\n'
         '  - Friuli-Venezia Giulia: Trieste - Udine - Pordenone - Gorizia\n'
         '  - Campania: Napoli - Caserta - Salerno - Benevento'),

        ('Independent Cities',
         '3 out of 107 cities showed no significant connections with other cities '
         '(R-squared < 0.3). Possible explanations include unique geographic characteristics '
         '(e.g., island locations), different public health policies, or distinct behavioral '
         'patterns in those regions.'),

        ('Advantages of CGP over Traditional Methods',
         'Key advantages of using CGP compared to simple linear correlation:\n\n'
         '  - CGP can identify non-linear relationships between cities\n'
         '  - Only truly influential inputs (Active Inputs) remain in the final model\n'
         '  - Temporal delays (Lags) are incorporated into the model\n'
         '  - The generated mathematical expressions are interpretable'),
    ]

    for title, content in sections:
        doc.add_heading(title, level=3)
        doc.add_paragraph(content)
        doc.add_paragraph('')
    doc.add_page_break()

    # === 7. CONCLUSIONS ===
    doc.add_heading('7. Conclusions', level=1)
    doc.add_paragraph(
        'In this study, Cartesian Genetic Programming (CGP) was used to investigate direct '
        'relationships between 107 Italian provinces regarding COVID-19 spread. The main findings are:'
    )

    conclusions = [
        f'Out of {n_total} connection pairs analyzed, {n_significant} significant connections (R-squared > 0.3) were identified',
        f'{n_connected} out of 107 cities had at least one significant direct connection',
        'The strongest connections were observed between geographically adjacent cities (R-squared up to 0.9998)',
        'Clear regional clusters were identified that align perfectly with Italy\'s geographic divisions',
        'Hub cities such as Arezzo, Trento, and Bari play central roles in the transmission network',
        'CGP successfully modeled non-linear relationships and temporal delays between cities',
        'These findings can inform regional public health policies and disease spread prediction models',
    ]
    for item in conclusions:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Future work should extend this analysis to data from 2020-2024 to study temporal '
        'changes in connection patterns. Additionally, combining CGP results with the SIR model '
        'could enable more comprehensive modeling of inter-city disease transmission dynamics.'
    )

    # === APPENDIX ===
    doc.add_page_break()
    doc.add_heading('Appendix: Complete List of Connections', level=1)

    all_headers = ['#', 'City 1', 'City 2', 'R-squared', 'Significant']
    all_rows = []
    for idx, row in conn_df.iterrows():
        all_rows.append([
            str(idx + 1), str(row['city_1']), str(row['city_2']),
            str(row['connection_strength_R2']),
            'Yes' if row['is_significant'] else 'No',
        ])
    add_styled_table(doc, all_headers, all_rows)

    output_path = os.path.join(OUTPUT_DIR, 'CGP_Report_English.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == '__main__':
    main()
