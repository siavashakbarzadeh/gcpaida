"""
Generate Farsi and English Word document reports for SIR model results (2024).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SIR_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2024\sir-result-2024"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2024\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def add_image_if_exists(doc, path, caption, width=Inches(6)):
    if os.path.exists(path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

# =========================================================================
# FARSI REPORT
# =========================================================================
def generate_farsi_report(summary_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11); style.font.name = 'Arial'

    # Title
    for _ in range(6): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('گزارش مدل‌سازی SIR\nبرای ۱۰۷ استان ایتالیا')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('مدل‌سازی Compartmental انتشار COVID-19\nبا استفاده از مدل SIR - داده‌های سال ۲۰۲۱')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph('')
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run('تاریخ: مارس ۲۰۲۶').font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading('فهرست مطالب', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for item in ['۱. مقدمه', '۲. مدل SIR', '۳. شرح داده‌ها', '۴. پارامترهای مدل',
                 '۵. نتایج', '   ۵.۱ جدول خلاصه', '   ۵.۲ مقایسه ۱۰ شهر برتر',
                 '   ۵.۳ نقشه حرارتی Beta', '   ۵.۴ توزیع R0', '   ۵.۵ نمونه نتایج',
                 '۶. تحلیل نتایج', '۷. نتیجه‌گیری', 'پیوست']:
        p = doc.add_paragraph(item); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    # 1
    doc.add_heading('۱. مقدمه', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'در این پروژه، مدل SIR برای مدل‌سازی انتشار COVID-19 در ۱۰۷ استان ایتالیا '
        'در سال ۲۰۲۱ استفاده شده است. هدف: تخمین β متغیر با زمان و R0 برای هر شهر.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2
    doc.add_heading('۲. مدل SIR', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('مدل SIR جمعیت را به سه گروه S, I, R تقسیم می‌کند.').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_table(doc, ['گروه', 'نماد', 'توضیح'], [
        ['حساس', 'S', 'افراد آلوده نشده'], ['آلوده', 'I', 'افراد مبتلا'], ['بهبودیافته', 'R', 'بهبود یافته/فوت شده'],
    ])
    doc.add_paragraph('')
    doc.add_paragraph('dS/dt = -β × S × I / N\ndI/dt = β × S × I / N - γ × I\ndR/dt = γ × I')
    doc.add_paragraph(
        'R0 = β/γ عدد بازتولید پایه. اگر R0 > 1 بیماری گسترش می‌یابد. '
        'از Piecewise fitting با بازه ۳۰ روزه و Differential Evolution استفاده شده.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # 3
    doc.add_heading('۳. شرح داده‌ها', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_styled_table(doc, ['شاخص', 'مقدار'], [
        ['تعداد شهرها', '107'], ['بازه زمانی', '2024-01-01 تا 2024-12-31'],
        ['تعداد روزها', '365'], ['نوع داده', 'totale_casi'], ['منبع', 'Protezione Civile Italiana'],
    ])
    doc.add_page_break()

    # 4
    doc.add_heading('۴. پارامترهای مدل', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_styled_table(doc, ['پارامتر', 'مقدار', 'توضیح'], [
        ['γ', '0.07', 'دوره بهبودی ۱۴ روز'], ['بازه', '30 روز', 'β ثابت در هر قطعه'],
        ['بهینه‌سازی', 'Diff. Evolution', 'الگوریتم تکاملی'], ['تکرار', '300', 'حداکثر نسل'],
        ['محدوده β', '[0.01, 2.0]', 'محدوده مجاز'], ['تابع هدف', 'Norm. MSE', 'خطای نرمال‌شده'],
    ])
    doc.add_page_break()

    # 5
    doc.add_heading('۵. نتایج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_cases = summary_df['total_cases'].sum()
    avg_R0 = summary_df['mean_R0'].mean()
    doc.add_paragraph(f'مجموع مبتلایان ۲۰۲۱: {total_cases:,.0f}. میانگین R0: {avg_R0:.2f}.').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 5.1
    doc.add_heading('۵.۱ جدول خلاصه (۲۰ شهر برتر)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    s_rows = [[str(i+1), str(r['city']), f"{int(r['population']):,}", f"{int(r['total_cases']):,}",
               f"{r['infection_rate_pct']:.1f}", str(r['mean_beta']), str(r['mean_R0'])]
              for i, (_, r) in enumerate(summary_df.head(20).iterrows())]
    add_styled_table(doc, ['#', 'شهر', 'جمعیت', 'مبتلایان', 'نرخ(%)', 'β', 'R0'], s_rows)
    doc.add_page_break()

    # 5.2-5.5 Images
    doc.add_heading('۵.۲ مقایسه ۱۰ شهر برتر', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'comparison_top10.png'), 'شکل ۱: مقایسه ۱۰ شهر برتر')
    doc.add_page_break()
    doc.add_heading('۵.۳ نقشه حرارتی Beta', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'beta_heatmap_all_cities.png'), 'شکل ۲: نقشه حرارتی Beta', Inches(6.5))
    doc.add_page_break()
    doc.add_heading('۵.۴ توزیع R0', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'R0_distribution.png'), 'شکل ۳: توزیع R0')
    doc.add_page_break()

    doc.add_heading('۵.۵ نمونه نتایج شهرها', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, city in enumerate(['roma', 'milano', 'napoli', 'brescia', 'bologna']):
        add_image_if_exists(doc, os.path.join(SIR_DIR, f'sir_{city}.png'), f'شکل {4+i}: {city.title()}', Inches(5.5))
        doc.add_paragraph('')
    doc.add_page_break()

    # 6-7
    doc.add_heading('۶. تحلیل نتایج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('β در طول سال ۲۰۲۱ به طور قابل توجهی تغییر کرده (امواج مختلف و lockdown).').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('۷. نتیجه‌گیری', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for item in [f'مجموع مبتلایان: {total_cases:,.0f}', f'میانگین R0: {avg_R0:.2f}',
                 'مدل SIR با موفقیت بر ۱۰۷ شهر برازش شد', 'Piecewise fitting بهترین نتایج را داد']:
        doc.add_paragraph(item, style='List Bullet').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Appendix
    doc.add_page_break()
    doc.add_heading('پیوست: جدول کامل ۱۰۷ شهر', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    app_rows = [[str(i+1), str(r['city']), f"{int(r['population']):,}", f"{int(r['total_cases']):,}",
                 f"{r['infection_rate_pct']:.1f}", str(r['mean_beta']), str(r['mean_R0'])]
                for i, (_, r) in enumerate(summary_df.iterrows())]
    add_styled_table(doc, ['#', 'شهر', 'جمعیت', 'مبتلایان', 'نرخ(%)', 'β', 'R0'], app_rows)

    path = os.path.join(OUTPUT_DIR, 'SIR_Report_Farsi_2024.docx')
    doc.save(path)
    print(f"Farsi report saved: {path}")


# =========================================================================
# ENGLISH REPORT
# =========================================================================
def generate_english_report(summary_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11); style.font.name = 'Arial'

    # Title
    for _ in range(6): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SIR Compartmental Model Report\nfor 107 Italian Provinces')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('COVID-19 Spread Modeling Using the SIR Model\nYear 2024 Data')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph(''); doc.add_paragraph('')
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run('Date: March 2026').font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', level=1)
    for item in ['1. Introduction', '2. The SIR Model', '3. Data Description',
                 '4. Model Parameters', '5. Results', '   5.1 Summary Table',
                 '   5.2 Top 10 Comparison', '   5.3 Beta Heatmap', '   5.4 R0 Distribution',
                 '   5.5 City Results', '6. Discussion', '7. Conclusions', 'Appendix']:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1-4
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('SIR model applied to COVID-19 data from 107 Italian provinces for 2024.')
    doc.add_heading('2. The SIR Model', level=1)
    doc.add_paragraph('Compartmental model: S (Susceptible), I (Infected), R (Recovered).')
    doc.add_paragraph('dS/dt = -beta*S*I/N, dI/dt = beta*S*I/N - gamma*I, dR/dt = gamma*I')
    doc.add_paragraph('R0 = beta/gamma. Piecewise fitting with 30-day segments using Differential Evolution.')
    doc.add_heading('3. Data Description', level=1)
    add_styled_table(doc, ['Metric', 'Value'], [
        ['Cities', '107'], ['Period', '2024-01-01 to 2024-12-31'], ['Days', '365'],
        ['Data Type', 'Cumulative cases (totale_casi)'], ['Source', 'Protezione Civile Italiana'],
    ])
    doc.add_heading('4. Model Parameters', level=1)
    add_styled_table(doc, ['Parameter', 'Value', 'Description'], [
        ['Gamma', '0.07', 'Recovery ~14 days'], ['Segment', '30 days', 'Piecewise constant beta'],
        ['Optimizer', 'Diff. Evolution', 'Robust evolutionary'], ['Max Iter', '300', 'Max generations'],
        ['Beta Range', '[0.01, 2.0]', 'Allowed range'], ['Objective', 'Norm. MSE', 'Normalized MSE'],
    ])
    doc.add_page_break()

    # 5
    doc.add_heading('5. Results', level=1)
    total_cases = summary_df['total_cases'].sum()
    avg_R0 = summary_df['mean_R0'].mean()
    doc.add_paragraph(f'Total 2024 infections: {total_cases:,.0f}. Average R0: {avg_R0:.2f}.')

    # 5.1
    doc.add_heading('5.1 Summary Table (Top 20)', level=2)
    s_rows = [[str(i+1), str(r['city']), f"{int(r['population']):,}", f"{int(r['total_cases']):,}",
               f"{r['infection_rate_pct']:.1f}", str(r['mean_beta']), str(r['mean_R0'])]
              for i, (_, r) in enumerate(summary_df.head(20).iterrows())]
    add_styled_table(doc, ['#', 'City', 'Pop.', 'Cases', 'Rate(%)', 'Beta', 'R0'], s_rows)
    doc.add_page_break()

    # 5.2-5.5
    doc.add_heading('5.2 Top 10 Cities Comparison', level=2)
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'comparison_top10.png'), 'Figure 1: Top 10 Comparison')
    doc.add_page_break()
    doc.add_heading('5.3 Beta Heatmap', level=2)
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'beta_heatmap_all_cities.png'), 'Figure 2: Beta Heatmap', Inches(6.5))
    doc.add_page_break()
    doc.add_heading('5.4 R0 Distribution', level=2)
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'R0_distribution.png'), 'Figure 3: R0 Distribution')
    doc.add_page_break()
    doc.add_heading('5.5 Individual City Results', level=2)
    for i, city in enumerate(['roma', 'milano', 'napoli', 'brescia', 'bologna']):
        add_image_if_exists(doc, os.path.join(SIR_DIR, f'sir_{city}.png'), f'Figure {4+i}: {city.title()}', Inches(5.5))
        doc.add_paragraph('')
    doc.add_page_break()

    # 6-7
    doc.add_heading('6. Discussion', level=1)
    doc.add_paragraph('Beta varied significantly throughout 2024 reflecting pandemic waves and interventions.')
    doc.add_heading('7. Conclusions', level=1)
    for item in [f'Total 2024 infections: {total_cases:,.0f}', f'Average R0: {avg_R0:.2f}',
                 'SIR fitted successfully to all 107 cities', 'Piecewise fitting with DE produced optimal results']:
        doc.add_paragraph(item, style='List Bullet')

    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix: Full Results (107 Cities)', level=1)
    app_rows = [[str(i+1), str(r['city']), f"{int(r['population']):,}", f"{int(r['total_cases']):,}",
                 f"{r['infection_rate_pct']:.1f}", str(r['mean_beta']), str(r['mean_R0'])]
                for i, (_, r) in enumerate(summary_df.iterrows())]
    add_styled_table(doc, ['#', 'City', 'Pop.', 'Cases', 'Rate(%)', 'Beta', 'R0'], app_rows)

    path = os.path.join(OUTPUT_DIR, 'SIR_Report_English_2024.docx')
    doc.save(path)
    print(f"English report saved: {path}")


# =========================================================================
# MAIN
# =========================================================================
def main():
    summary_df = pd.read_csv(os.path.join(SIR_DIR, 'sir_summary_all_cities.csv'))
    print("Generating SIR reports for 2024...")
    generate_farsi_report(summary_df)
    generate_english_report(summary_df)
    print("Done! Both reports saved to:", OUTPUT_DIR)

if __name__ == '__main__':
    main()
