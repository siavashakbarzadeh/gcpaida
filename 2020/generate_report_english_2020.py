"""
Generate an English Word document report for CGP city analysis results (2020).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

CGP_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020\cgp_results"
OUTPUT_DIR = r"c:\Users\Utente\Documents\GitHub\gcpaida\2020\reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11); style.font.name = 'Arial'

    # Title
    for _ in range(6): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Inter-City Infection Relationship Analysis\nUsing Cartesian Genetic Programming (CGP)')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Behavioral Analysis of COVID-19 Spread Across 107 Italian Provinces\nYear 2020 Data')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph(''); doc.add_paragraph('')
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run('Date: March 2026').font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', level=1)
    for item in ['1. Introduction and Objectives', '2. CGP Overview', '3. Data Description',
                 '4. CGP Model Parameters', '5. CGP Analysis Results',
                 '   5.1 Strongest Connections', '   5.2 Hub Cities', '   5.3 Network Graph',
                 '   5.4 Connection Heatmap', '   5.5 Top 30 Matrix',
                 '   5.6 Fitness Convergence', '   5.7 Hub Influence',
                 '6. Discussion', '7. Conclusions', 'Appendix']:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Sections 1-4
    doc.add_heading('1. Introduction and Objectives', level=1)
    doc.add_paragraph(
        'This study investigates direct relationships between Italian cities regarding COVID-19 spread '
        'using CGP on 2020 data from 107 provinces.')
    doc.add_heading('2. CGP Overview', level=1)
    doc.add_paragraph('CGP is an evolutionary algorithm by Julian Miller (2000) using DAGs on a 2D grid.')
    doc.add_heading('3. Data Description', level=1)
    doc.add_paragraph('Daily cumulative infections for 107 provinces in 2020 (365 days), merged into merged_cities_2020.csv.')
    doc.add_heading('4. CGP Model Parameters', level=1)
    add_styled_table(doc, ['Parameter', 'Value', 'Description'], [
        ['Grid', '3x8', 'CGP grid dimensions'], ['Generations', '300', 'Max generations'],
        ['Lambda', '4', 'Offspring per gen'], ['Mutation Rate', '0.08', 'Per-gene mutation prob'],
        ['Lags', '3', '0,1,2 day lags'], ['Candidates', '15', 'Top correlated cities'],
        ['R2 Threshold', '0.3', 'Significance cutoff'], ['Functions', '8', 'Math operations'],
    ])
    doc.add_page_break()

    # 5. Results
    doc.add_heading('5. CGP Analysis Results', level=1)
    conn_df = pd.read_csv(os.path.join(CGP_DIR, 'all_connections.csv'))
    summary_df = pd.read_csv(os.path.join(CGP_DIR, 'city_summary.csv'))
    n_total = len(conn_df)
    n_sig = len(conn_df[conn_df['is_significant'] == True])
    n_conn = len(summary_df[summary_df['n_connections'] > 0])
    doc.add_paragraph(f'{n_total} pairs found, {n_sig} significant. {n_conn}/107 cities connected.')

    # 5.1
    doc.add_heading('5.1 Strongest Connections', level=2)
    top_h = ['Rank', 'City 1', 'City 2', 'R2', 'Sig']
    top_r = [[str(i+1), str(r['city_1']), str(r['city_2']), str(r['connection_strength_R2']),
              'Yes' if r['is_significant'] else 'No'] for i, (_, r) in enumerate(conn_df.head(20).iterrows())]
    add_styled_table(doc, top_h, top_r)
    img = os.path.join(CGP_DIR, 'top_connections.png')
    if os.path.exists(img):
        doc.add_paragraph(''); doc.add_picture(img, width=Inches(6))
    doc.add_page_break()

    # 5.2
    doc.add_heading('5.2 Hub Cities', level=2)
    hub_r = [[str(i+1), str(r['city']), str(r['n_connections']),
              str(r['linked_cities'])[:60]] for i, (_, r) in enumerate(summary_df.head(15).iterrows())]
    add_styled_table(doc, ['Rank', 'City', 'Connections', 'Linked'], hub_r)
    img = os.path.join(CGP_DIR, 'hub_cities.png')
    if os.path.exists(img):
        doc.add_paragraph(''); doc.add_picture(img, width=Inches(6))
    doc.add_page_break()

    # 5.3-5.7 Images
    for section, title, fname in [
        ('5.3', 'Network Graph', 'network_graph.png'),
        ('5.4', 'Connection Heatmap', 'connection_heatmap.png'),
        ('5.5', 'Top 30 Matrix', 'top30_connection_matrix.png'),
        ('5.6', 'Fitness Convergence', 'cgp_fitness_convergence.png'),
    ]:
        doc.add_heading(f'{section} {title}', level=2)
        img = os.path.join(CGP_DIR, fname)
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(6))
        doc.add_page_break()

    doc.add_heading('5.7 Hub Influence', level=2)
    for f in sorted([f for f in os.listdir(CGP_DIR) if f.startswith('influence_')]):
        img = os.path.join(CGP_DIR, f)
        doc.add_picture(img, width=Inches(4.5)); doc.add_paragraph('')
    doc.add_page_break()

    # 6-7
    doc.add_heading('6. Discussion', level=1)
    doc.add_paragraph('Geographic proximity drives connections. Clear regional clusters found.')
    doc.add_heading('7. Conclusions', level=1)
    doc.add_paragraph(f'{n_sig} significant connections found among 107 cities in 2020 data.')

    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix: All Connections', level=1)
    all_r = [[str(i+1), str(r['city_1']), str(r['city_2']), str(r['connection_strength_R2']),
              'Yes' if r['is_significant'] else 'No'] for i, (_, r) in enumerate(conn_df.iterrows())]
    add_styled_table(doc, ['#', 'City 1', 'City 2', 'R2', 'Sig'], all_r)

    path = os.path.join(OUTPUT_DIR, 'CGP_Report_English_2020.docx')
    doc.save(path)
    print(f"Report saved to: {path}")

if __name__ == '__main__':
    main()
