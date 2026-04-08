"""
Generate SIR Model Reports (Farsi + English) for Synthetic Data
================================================================
Produces Word documents summarizing the SIR model results.
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SIR_DIR = os.path.join(BASE_DIR, 'sir_results')
OUTPUT_DIR = os.path.join(BASE_DIR, 'reports')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ==============================================================================
# HELPERS
# ==============================================================================

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def add_image_if_exists(doc, path, caption, width=Inches(6)):
    if os.path.exists(path):
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True; r.font.size = Pt(10)
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
        return True
    return False


# ==============================================================================
# FARSI REPORT
# ==============================================================================

def generate_farsi_report(summary_df, gt_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11); style.font.name = 'Arial'

    # Title Page
    for _ in range(5): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('گزارش مدل SIR\nبرای داده‌های سینتتیک ۱۰ شهر ایتالیا')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('مدل‌سازی SIR با پارامترهای γ=۰.۱۴ و β=۰.۳۵\nداده‌های سینتتیک با اتصالات مشخص')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph('')
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run('تاریخ: آوریل ۲۰۲۶').font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading('فهرست مطالب', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for item in ['۱. مقدمه', '۲. مدل SIR', '۳. شرح داده‌های سینتتیک',
                 '۴. پارامترهای مدل', '۵. اتصالات شناخته‌شده',
                 '۶. نتایج', '۷. نتیجه‌گیری']:
        p = doc.add_paragraph(item); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('۱. مقدمه', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'در این پروژه، از مدل SIR (حساس-آلوده-بهبودیافته) برای تولید داده‌های مصنوعی (سینتتیک) '
        'COVID-19 در ۱۰ شهر ایتالیا استفاده شده است. هدف اصلی ساخت دادگانی با اتصالات بین‌شهری '
        'مشخص است تا در مرحله بعد بتوانیم با الگوریتم CGP این اتصالات را بازیابی و اعتبارسنجی کنیم.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        'مزیت استفاده از داده‌های سینتتیک این است که "حقیقت زمینی" (Ground Truth) را می‌دانیم '
        'و می‌توانیم عملکرد CGP را به صورت دقیق ارزیابی کنیم.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2. SIR Model
    doc.add_heading('۲. مدل SIR', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'مدل SIR جمعیت را به سه گروه تقسیم می‌کند:'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_styled_table(doc, ['گروه', 'نماد', 'توضیح'], [
        ['حساس', 'S', 'افرادی که هنوز آلوده نشده‌اند'],
        ['آلوده', 'I', 'افراد مبتلا که قادر به انتقال هستند'],
        ['بهبودیافته', 'R', 'افراد بهبودیافته یا فوت‌شده'],
    ])
    doc.add_paragraph('')
    doc.add_paragraph(
        'معادلات دیفرانسیل:\n'
        'dS/dt = -β × S × I / N\n'
        'dI/dt = β × S × I / N - γ × I\n'
        'dR/dt = γ × I'
    )
    doc.add_paragraph(
        'R₀ = β/γ عدد بازتولید پایه است. اگر R₀ > 1 بیماری گسترش می‌یابد.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # 3. Data Description
    doc.add_heading('۳. شرح داده‌های سینتتیک', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'داده‌های سینتتیک برای ۱۰ شهر با مدل SIR تولید شده‌اند. هر شهر شرایط اولیه متفاوتی دارد '
        'و بین شهرها اثر سرایت (Spillover) با درصدهای مختلف (۱٪، ۵٪، ۱۰٪) اعمال شده است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # City table
    city_headers = ['#', 'شهر', 'منطقه', 'جمعیت', 'مبتلایان', 'نرخ(٪)', 'R₀']
    city_rows = []
    for i, (_, row) in enumerate(summary_df.iterrows()):
        city_rows.append([
            str(i+1), str(row['city']), str(row['region']),
            f"{int(row['population']):,}", f"{int(row['total_cases']):,}",
            f"{row['infection_rate_pct']:.1f}", str(row['mean_R0'])
        ])
    add_styled_table(doc, city_headers, city_rows)
    doc.add_page_break()

    # 4. Parameters
    doc.add_heading('۴. پارامترهای مدل', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_styled_table(doc, ['پارامتر', 'مقدار', 'توضیح'], [
        ['γ (گاما)', '0.14', 'نرخ بهبودی: دوره بهبودی ≈ ۷ روز'],
        ['β (بتا)', '0.35', 'نرخ آلودگی'],
        ['R₀', '2.5', 'عدد بازتولید پایه (β/γ)'],
        ['مدت شبیه‌سازی', '365 روز', 'یک سال کامل'],
        ['تعداد شهرها', '10', 'شهرهای اصلی ایتالیا'],
        ['سرایت بین‌شهری', '1%, 5%, 10%', 'سه سطح مختلف'],
    ])
    doc.add_page_break()

    # 5. Ground Truth Connections
    doc.add_heading('۵. اتصالات شناخته‌شده (Ground Truth)', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'ارتباطات زیر به عنوان حقیقت زمینی تعریف شده‌اند. '
        'هدف CGP بازیابی این ارتباطات از داده‌های سینتتیک است:'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    gt_headers = ['#', 'شهر مبدا', 'شهر مقصد', 'قدرت', 'تاخیر (روز)']
    gt_rows = []
    for i, (_, row) in enumerate(gt_df.iterrows()):
        gt_rows.append([
            str(i+1), str(row['source']), str(row['target']),
            'قوی' if row['strength'] == 'strong' else 'متوسط',
            str(row['lag_days'])
        ])
    add_styled_table(doc, gt_headers, gt_rows)
    doc.add_paragraph('')

    add_image_if_exists(doc, os.path.join(SIR_DIR, 'ground_truth_network.png'),
                        'شکل ۱: گراف شبکه اتصالات شناخته‌شده', Inches(5.5))
    doc.add_page_break()

    # 6. Results
    doc.add_heading('۶. نتایج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('۶.۱ مقایسه همه شهرها', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'comparison_all_cities.png'),
                        'شکل ۲: مقایسه ده شهر - مدل SIR سینتتیک')
    doc.add_page_break()

    doc.add_heading('۶.۲ اثر درصد سرایت', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'spillover_comparison.png'),
                        'شکل ۳: مقایسه اثر ۱٪ و ۵٪ و ۱۰٪ سرایت')
    doc.add_page_break()

    doc.add_heading('۶.۳ نتایج شهرهای نمونه', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, city in enumerate(['milano', 'roma', 'napoli', 'torino', 'bologna']):
        add_image_if_exists(doc, os.path.join(SIR_DIR, f'sir_{city}.png'),
                            f'شکل {4+i}: نتایج SIR - {city.title()}', Inches(5.5))
    doc.add_page_break()

    # 7. Conclusion
    doc.add_heading('۷. نتیجه‌گیری', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    conclusions = [
        f'مجموع ۱۰ شهر با جمعیت‌های واقعی شبیه‌سازی شد',
        f'مدل SIR با β=0.35 و γ=0.14 (R₀=2.5) اجرا شد',
        'اتصالات بین‌شهری با سه سطح سرایت (۱٪، ۵٪، ۱۰٪) اعمال شد',
        'داده‌ها آماده تحلیل CGP برای بازیابی اتصالات هستند',
        'حقیقت زمینی ۱۰ ارتباط مشخص بین شهرها تعریف شده است',
    ]
    for item in conclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(OUTPUT_DIR, 'SIR_Report_Farsi_SynData.docx')
    doc.save(path)
    print(f"  Farsi SIR report saved: {path}")


# ==============================================================================
# ENGLISH REPORT
# ==============================================================================

def generate_english_report(summary_df, gt_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11); style.font.name = 'Times New Roman'

    # Title Page
    for _ in range(5): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SIR Model Report\nSynthetic Data for 10 Italian Cities')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(46, 64, 87)
    doc.add_paragraph('')
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('SIR Modeling with γ=0.14 and β=0.35\n'
                   'Synthetic Data with Known Inter-City Connections')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph('')
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run('Date: April 2026').font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', level=1)
    for item in ['1. Introduction', '2. The SIR Model', '3. Synthetic Data Description',
                 '4. Model Parameters', '5. Ground Truth Connections',
                 '6. Results', '7. Conclusions']:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This project uses the SIR (Susceptible-Infected-Recovered) compartmental model '
        'to generate synthetic COVID-19 infection data for 10 Italian cities. The primary '
        'goal is to create a dataset with KNOWN inter-city connections (ground truth) so '
        'that the CGP algorithm can be validated by recovering these connections.'
    )
    doc.add_paragraph(
        'The advantage of synthetic data is that we know the "ground truth" and can '
        'precisely evaluate CGP\'s ability to discover the correct connections.'
    )

    # 2. SIR Model
    doc.add_heading('2. The SIR Model', level=1)
    doc.add_paragraph(
        'The SIR model divides the population into three compartments:'
    )
    add_styled_table(doc, ['Compartment', 'Symbol', 'Description'], [
        ['Susceptible', 'S', 'Individuals not yet infected'],
        ['Infected', 'I', 'Currently infectious individuals'],
        ['Recovered', 'R', 'Recovered or deceased individuals'],
    ])
    doc.add_paragraph('')
    doc.add_paragraph(
        'Differential equations:\n'
        'dS/dt = -β × S × I / N\n'
        'dI/dt = β × S × I / N - γ × I\n'
        'dR/dt = γ × I\n\n'
        'R₀ = β/γ is the basic reproduction number. When R₀ > 1, the disease spreads.'
    )
    doc.add_page_break()

    # 3. Data Description
    doc.add_heading('3. Synthetic Data Description', level=1)
    doc.add_paragraph(
        'Synthetic data was generated for 10 Italian cities using the SIR model '
        'with city-specific initial conditions and population sizes. Inter-city '
        'spillover effects were applied at three levels (1%, 5%, 10%) to simulate '
        'disease transmission between connected cities.'
    )

    city_headers = ['#', 'City', 'Region', 'Population', 'Total Cases',
                    'Rate(%)', 'R₀']
    city_rows = []
    for i, (_, row) in enumerate(summary_df.iterrows()):
        city_rows.append([
            str(i+1), str(row['city']), str(row['region']),
            f"{int(row['population']):,}", f"{int(row['total_cases']):,}",
            f"{row['infection_rate_pct']:.1f}", str(row['mean_R0'])
        ])
    add_styled_table(doc, city_headers, city_rows)
    doc.add_page_break()

    # 4. Parameters
    doc.add_heading('4. Model Parameters', level=1)
    add_styled_table(doc, ['Parameter', 'Value', 'Description'], [
        ['γ (gamma)', '0.14', 'Recovery rate: ~7 day recovery period'],
        ['β (beta)', '0.35', 'Infection rate'],
        ['R₀', '2.5', 'Basic reproduction number (β/γ)'],
        ['Duration', '365 days', 'Full year simulation'],
        ['Cities', '10', 'Major Italian cities'],
        ['Spillover', '1%, 5%, 10%', 'Three levels tested'],
    ])
    doc.add_page_break()

    # 5. Ground Truth
    doc.add_heading('5. Ground Truth Connections', level=1)
    doc.add_paragraph(
        'The following connections are defined as ground truth. '
        'The goal of CGP is to recover these connections from the synthetic data:'
    )

    gt_headers = ['#', 'Source', 'Target', 'Strength', 'Lag (days)']
    gt_rows = []
    for i, (_, row) in enumerate(gt_df.iterrows()):
        gt_rows.append([
            str(i+1), str(row['source']), str(row['target']),
            str(row['strength']).title(), str(row['lag_days'])
        ])
    add_styled_table(doc, gt_headers, gt_rows)
    doc.add_paragraph('')

    add_image_if_exists(doc, os.path.join(SIR_DIR, 'ground_truth_network.png'),
                        'Figure 1: Ground Truth Network of City Connections',
                        Inches(5.5))
    doc.add_page_break()

    # 6. Results
    doc.add_heading('6. Results', level=1)

    doc.add_heading('6.1 All Cities Comparison', level=2)
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'comparison_all_cities.png'),
                        'Figure 2: Comparison of All 10 Cities - Synthetic SIR')
    doc.add_page_break()

    doc.add_heading('6.2 Spillover Effect', level=2)
    add_image_if_exists(doc, os.path.join(SIR_DIR, 'spillover_comparison.png'),
                        'Figure 3: Comparing 1%, 5%, and 10% Spillover Effects')
    doc.add_page_break()

    doc.add_heading('6.3 Sample City Results', level=2)
    for i, city in enumerate(['milano', 'roma', 'napoli', 'torino', 'bologna']):
        add_image_if_exists(doc, os.path.join(SIR_DIR, f'sir_{city}.png'),
                            f'Figure {4+i}: SIR Results - {city.title()}',
                            Inches(5.5))
    doc.add_page_break()

    # 7. Conclusions
    doc.add_heading('7. Conclusions', level=1)
    conclusions = [
        '10 Italian cities simulated with real population data',
        'SIR model with β=0.35, γ=0.14 (R₀=2.5) produced realistic curves',
        'Inter-city spillover applied at three levels (1%, 5%, 10%)',
        'Data is ready for CGP analysis to validate connection discovery',
        'Ground truth defines 10 specific city-to-city connections',
    ]
    for item in conclusions:
        doc.add_paragraph(item, style='List Bullet')

    path = os.path.join(OUTPUT_DIR, 'SIR_Report_English_SynData.docx')
    doc.save(path)
    print(f"  English SIR report saved: {path}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    print("Generating SIR Reports for Synthetic Data...")

    summary_path = os.path.join(SIR_DIR, 'sir_summary.csv')
    gt_path = os.path.join(SIR_DIR, 'ground_truth_connections.csv')

    if not os.path.exists(summary_path):
        print(f"ERROR: {summary_path} not found! Run syn_data_sir_model.py first.")
        return
    if not os.path.exists(gt_path):
        print(f"ERROR: {gt_path} not found! Run syn_data_sir_model.py first.")
        return

    summary_df = pd.read_csv(summary_path)
    gt_df = pd.read_csv(gt_path)

    generate_farsi_report(summary_df, gt_df)
    generate_english_report(summary_df, gt_df)

    print(f"\nDone! Reports saved to: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
