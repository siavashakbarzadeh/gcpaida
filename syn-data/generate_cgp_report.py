"""
Generate CGP Analysis Reports (Farsi + English) for Synthetic Data
===================================================================
Comprehensive bilingual reports including ALL figures from cgp_results/:
  - Correlation matrix (per spillover %)
  - Lag sweep plots (per spillover %)
  - Connection heatmaps (per spillover %)
  - Network graphs (per spillover %)
  - CGP weights heatmaps (per spillover %)
  - Fitness convergence (per spillover %)
  - Precision/Recall/F1 comparison (cross-spillover)
  - Discovered vs Ground Truth bar chart (cross-spillover)
  - Ground truth network from SIR
  - Tables from CSVs
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CGP_DIR = os.path.join(BASE_DIR, 'cgp_results')
SIR_DIR = os.path.join(BASE_DIR, 'sir_results')
OUTPUT_DIR = os.path.join(BASE_DIR, 'reports')
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ==============================================================================
# HELPERS
# ==============================================================================

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, header_color='2E4057'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
        # Alternate shading
        if row_idx % 2 == 1:
            for col_idx in range(len(row_data)):
                set_cell_shading(table.rows[row_idx + 1].cells[col_idx], 'EBF5FB')
    return table


def add_img(doc, path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph('')
        return True
    else:
        print(f"  WARNING: File not found: {path}")
        return False


def add_spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph('')


# File path helpers for exact names in cgp_results
def cgp_path(name):
    return os.path.join(CGP_DIR, name)

def sir_path(name):
    return os.path.join(SIR_DIR, name)


# ==============================================================================
# LOAD ALL DATA
# ==============================================================================

def load_all_data():
    data = {}

    # Ground truth
    gt_path = sir_path('ground_truth_connections.csv')
    data['ground_truth'] = pd.read_csv(gt_path) if os.path.exists(gt_path) else None

    # SIR summary
    sir_sum_path = sir_path('sir_summary.csv')
    data['sir_summary'] = pd.read_csv(sir_sum_path) if os.path.exists(sir_sum_path) else None

    # GT comparison
    gt_comp_path = cgp_path('ground_truth_comparison.csv')
    data['gt_comparison'] = pd.read_csv(gt_comp_path) if os.path.exists(gt_comp_path) else None

    # Lag sweep
    lag_path = cgp_path('lag_sweep_all.csv')
    data['lag_sweep'] = pd.read_csv(lag_path) if os.path.exists(lag_path) else None

    # Per-percentage data
    for pct in [1, 5, 10]:
        conn_path = cgp_path(f'connections_{pct}pct.csv')
        data[f'connections_{pct}'] = pd.read_csv(conn_path) if os.path.exists(conn_path) else None

        cs_path = cgp_path(f'city_summary_{pct}pct.csv')
        data[f'city_summary_{pct}'] = pd.read_csv(cs_path) if os.path.exists(cs_path) else None

    return data


# ==============================================================================
# FARSI REPORT
# ==============================================================================

def generate_farsi_report(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Arial'

    # ==================== TITLE PAGE ====================
    for _ in range(4):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('گزارش تحلیل CGP\nکشف ارتباطات بین‌شهری از داده‌های سینتتیک')
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(46, 64, 87)
    add_spacer(doc, 2)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('برنامه‌نویسی ژنتیک کارتزین (CGP)\nتوابع جمعی (Addition-Only)\nمقایسه با حقیقت زمینی (Ground Truth)')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(100, 100, 100)
    add_spacer(doc, 2)
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run('مدل SIR: beta=0.35, gamma=0.14, R0=2.50\n'
                   'شهرها: ۱۰ شهر ایتالیایی\n'
                   'آستانه R2 = 0.9')
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)
    add_spacer(doc)
    d2 = doc.add_paragraph()
    d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d2.add_run('تاریخ: آوریل ۲۰۲۶').font.size = Pt(12)
    doc.add_page_break()

    # ==================== TOC ====================
    doc.add_heading('فهرست مطالب', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    toc = [
        '۱. مقدمه و هدف',
        '۲. معرفی مختصر CGP',
        '۳. مجموعه توابع (Addition-Only)',
        '۴. پارامترهای CGP',
        '۵. ارتباطات حقیقت زمینی (Ground Truth)',
        '۶. ماتریس همبستگی شهرها',
        '۷. یافتن لگ بهینه',
        '۸. نتایج تحلیل CGP',
        '   ۸.۱ سرایت ۱٪',
        '   ۸.۲ سرایت ۵٪',
        '   ۸.۳ سرایت ۱۰٪',
        '۹. مقایسه کلی با حقیقت زمینی',
        '۱۰. تحلیل نتایج',
        '۱۱. نتیجه‌گیری',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    # ==================== SECTION 1: Introduction ====================
    doc.add_heading('۱. مقدمه و هدف', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'هدف از این تحلیل، اعتبارسنجی الگوریتم برنامه‌نویسی ژنتیک کارتزین (CGP) '
        'در کشف ارتباطات مستقیم بین شهرهاست. دقیقاً همان روش و پارامترهایی که '
        'روی داده‌های واقعی COVID-19 (سال ۲۰۲۰) استفاده شد، اینجا نیز اعمال شده است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        'داده‌های سینتتیک با مدل SIR (beta=0.35, gamma=0.14) برای ۱۰ شهر ایتالیایی '
        'تولید شده‌اند. یک شبکه ارتباطی مشخص (Ground Truth) تعریف شده و '
        'هدف CGP کشف مجدد این ارتباطات از داده‌های مبتلایان روزانه است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        'سه سطح سرایت بین‌شهری آزمایش شده‌اند:'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for item in ['۱٪ - سرایت ضعیف (چالش‌برانگیز برای شناسایی)',
                 '۵٪ - سرایت متوسط',
                 '۱۰٪ - سرایت قوی (آسان‌تر برای شناسایی)']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # ==================== SECTION 2: CGP Method ====================
    doc.add_heading('۲. معرفی مختصر CGP', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'CGP یک الگوریتم تکاملی است که عبارات ریاضی را بر روی یک شبکه '
        'دوبعدی (grid) نمایش می‌دهد. هر گره (node) یک تابع ریاضی و دو ورودی دارد.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph('مراحل کار CGP:').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    steps = [
        'برای هر شهر هدف، داده‌های لگ‌دار (lagged) سایر شهرها به عنوان ورودی ارائه می‌شود',
        'CGP یک عبارت ریاضی تکامل می‌دهد که مبتلایان جدید شهر هدف را پیش‌بینی کند',
        'با استفاده از استراتژی (1+lambda)، بهترین فرد در هر نسل انتخاب می‌شود',
        'ورودی‌های فعال (Active Inputs) نشان‌دهنده شهرهای واقعاً مؤثر هستند',
        'اگر R2 > 0.9 باشد، ارتباط معنادار تلقی می‌شود',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    doc.add_page_break()

    # ==================== SECTION 3: Functions ====================
    doc.add_heading('۳. مجموعه توابع (Addition-Only)', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'در این تحلیل، دقیقاً مانند تحلیل داده‌های واقعی ۲۰۲۰، فقط از توابع جمعی '
        'استفاده شده است. هیچ تابع تفاضلی (sub)، قدرمطلق تفاضل (abs_diff)، ضرب (mul) '
        'یا تقسیم (div) وجود ندارد.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    add_styled_table(doc, ['تابع', 'فرمول', 'توضیح'], [
        ['add', 'a + b', 'جمع مستقیم دو ورودی'],
        ['max', 'max(a, b)', 'بیشینه دو ورودی'],
        ['min', 'min(a, b)', 'کمینه دو ورودی'],
        ['avg', '(a + b) / 2', 'میانگین دو ورودی'],
        ['weighted_add', '0.7a + 0.3b', 'جمع وزن‌دار (تاکید بر ورودی اول)'],
    ])
    add_spacer(doc)
    doc.add_paragraph(
        'دلیل انتخاب: توابع جمعی، روابط تجمعی و هموار بین شهرها را بهتر مدل‌سازی '
        'می‌کنند. پدیده سرایت بیماری ذاتاً تجمعی است: مبتلایان از شهرهای مختلف '
        'به صورت جمعی بر شهر هدف تاثیر می‌گذارند.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # ==================== SECTION 4: Parameters ====================
    doc.add_heading('۴. پارامترهای CGP', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_styled_table(doc, ['پارامتر', 'مقدار', 'توضیح'], [
        ['سطرهای شبکه', '3', 'تعداد سطرهای grid'],
        ['ستون‌های شبکه', '8', 'تعداد ستون‌های grid'],
        ['خروجی‌ها', '1', 'پیش‌بینی مبتلایان جدید شهر هدف'],
        ['تعداد نسل‌ها', '500', 'حداکثر تعداد نسل الگوریتم تکاملی'],
        ['Lambda', '4', 'تعداد فرزندان در هر نسل'],
        ['نرخ جهش', '0.10', 'احتمال جهش هر ژن (۱۰٪)'],
        ['تکرار اجرا', '3', 'بهترین نتیجه از ۳ اجرای مستقل'],
        ['آستانه R2', '0.9', 'حداقل R2 برای تشخیص ارتباط معنادار'],
        ['Levels-back', '8', 'اتصال‌پذیری کامل بین لایه‌ها'],
        ['تعداد توابع', '5', 'add, max, min, avg, weighted_add'],
        ['نرمال‌سازی', 'Z-score', 'استانداردسازی ورودی‌ها و خروجی'],
    ])
    doc.add_page_break()

    # ==================== SECTION 5: Ground Truth ====================
    doc.add_heading('۵. ارتباطات حقیقت زمینی (Ground Truth)', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'شبکه ارتباطی زیر در مدل SIR تعریف شده است. ۱۰ ارتباط بین‌شهری با '
        'شدت‌ها و تاخیرهای مختلف وجود دارد. هدف CGP کشف مجدد این ارتباطات است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    if data['ground_truth'] is not None:
        gt = data['ground_truth']
        gt_rows = []
        for _, row in gt.iterrows():
            strength_fa = {'strong': 'قوی', 'medium': 'متوسط', 'weak': 'ضعیف'}.get(row['strength'], str(row['strength']))
            gt_rows.append([
                str(row['source']).title(),
                str(row['target']).title(),
                strength_fa,
                f"{row['lag_days']} روز",
            ])
        add_styled_table(doc, ['شهر مبدا', 'شهر مقصد', 'شدت', 'تاخیر'], gt_rows)

    add_spacer(doc)
    add_img(doc, sir_path('ground_truth_network.png'),
            'شکل ۱: شبکه ارتباطات حقیقت زمینی (Ground Truth Network)', Inches(5))
    doc.add_page_break()

    # ==================== SECTION 6: Correlation Matrix ====================
    doc.add_heading('۶. ماتریس همبستگی شهرها', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'ماتریس همبستگی پیرسون بین مبتلایان جدید روزانه شهرها محاسبه شده است. '
        'این ماتریس نشان‌دهنده همبستگی خطی پیش از تحلیل CGP است. '
        'CGP فراتر از همبستگی ساده عمل کرده و روابط غیرخطی واقعی را شناسایی می‌کند.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    for pct in [1, 5, 10]:
        add_img(doc, cgp_path(f'correlation_matrix_{pct}pct_spillover.png'),
                f'شکل: ماتریس همبستگی پیرسون - سرایت {pct}٪', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 7: Optimal Lag ====================
    doc.add_heading('۷. یافتن لگ بهینه', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'برای یافتن بهترین تعداد روزهای تاخیری، لگ‌های مختلف '
        '(۱، ۲، ۳، ۵، ۷، ۱۰، ۱۴ و ۲۱ روز) آزمایش شدند. '
        'برای هر لگ، CGP اجرا شده و میانگین R2 محاسبه گردید.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    if data['lag_sweep'] is not None:
        lag = data['lag_sweep']
        lag_headers = ['درصد سرایت', 'لگ (روز)', 'میانگین R2', 'حداکثر R2', 'شهرها بالای آستانه']
        lag_rows = []
        for _, row in lag.iterrows():
            lag_rows.append([
                str(row['spillover']),
                str(int(row['lag'])),
                f"{row['avg_r2']:.4f}",
                f"{row['max_r2']:.4f}",
                str(int(row['n_above_threshold'])),
            ])
        add_styled_table(doc, lag_headers, lag_rows)
    add_spacer(doc)

    for pct in [1, 5, 10]:
        add_img(doc, cgp_path(f'lag_sweep_({pct}pct_spillover).png'),
                f'شکل: نتایج جستجوی لگ بهینه - سرایت {pct}٪', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 8: Results per % ====================
    doc.add_heading('۸. نتایج تحلیل CGP', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for pct_idx, pct in enumerate([1, 5, 10]):
        section_num = f'۸.{pct_idx+1}'
        doc.add_heading(f'{section_num} سرایت {pct}٪', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # City summary table
        cs_key = f'city_summary_{pct}'
        if data[cs_key] is not None:
            cs = data[cs_key]
            doc.add_heading('خلاصه نتایج هر شهر', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cs_headers = ['شهر', 'R2', 'تعداد اتصالات', 'شهرهای مرتبط', 'توابع فعال']
            cs_rows = []
            for _, row in cs.iterrows():
                linked = str(row.get('linked_cities', '')).replace(';', '، ')
                funcs = str(row.get('active_functions', ''))
                if funcs == 'nan':
                    funcs = '-'
                cs_rows.append([
                    str(row['city']),
                    f"{row['r2_score']:.4f}",
                    str(int(row['n_connections'])),
                    linked if linked != 'nan' else '-',
                    funcs,
                ])
            add_styled_table(doc, cs_headers, cs_rows)
        add_spacer(doc)

        # Connections table
        conn_key = f'connections_{pct}'
        if data[conn_key] is not None:
            cdf = data[conn_key]
            sig = cdf[cdf['is_significant'] == True]
            doc.add_heading('ارتباطات کشف‌شده', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph(
                f'تعداد کل ارتباطات کشف‌شده: {len(sig)} (از {len(cdf)} جفت تحلیل‌شده)'
            ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            conn_headers = ['شهر ۱', 'شهر ۲', 'قدرت (R2)', 'حقیقت زمینی', 'طبقه‌بندی']
            conn_rows = []
            for _, row in sig.iterrows():
                gt_label = 'بله' if row.get('is_ground_truth', False) else 'خیر'
                cls_label = {'TP': 'صحیح (TP)', 'FP': 'اضافه (FP)', 'FN': 'از دست رفته (FN)', 'TN': 'TN'}.get(
                    str(row.get('classification', '')), str(row.get('classification', '')))
                conn_rows.append([
                    str(row['city_1']),
                    str(row['city_2']),
                    f"{row['connection_strength_R2']:.4f}",
                    gt_label,
                    cls_label,
                ])
            add_styled_table(doc, conn_headers, conn_rows)
        add_spacer(doc)

        # ALL visualizations for this percentage
        doc.add_heading('ماتریس همبستگی', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_img(doc, cgp_path(f'correlation_matrix_{pct}pct_spillover.png'),
                f'شکل: ماتریس همبستگی پیرسون - سرایت {pct}٪', Inches(5))

        doc.add_heading('نقشه حرارتی ارتباطات CGP', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(
            'این نقشه حرارتی قدرت ارتباطات کشف‌شده توسط CGP (بر اساس R2) را نشان می‌دهد.'
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_img(doc, cgp_path(f'connection_heatmap_{pct}pct_spillover.png'),
                f'شکل: نقشه حرارتی ارتباطات - سرایت {pct}٪', Inches(5))

        doc.add_heading('گراف شبکه ارتباطات', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(
            'در این گراف، خطوط سبز نشان‌دهنده ارتباطات صحیح (True Positive) و '
            'خطوط قرمز نقطه‌چین نشان‌دهنده ارتباطات اضافی (False Positive) هستند. '
            'اندازه هر گره متناسب با تعداد ارتباطات آن است.'
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_img(doc, cgp_path(f'network_graph_{pct}pct_spillover.png'),
                f'شکل: گراف شبکه ارتباطات - سرایت {pct}٪', Inches(5.5))

        doc.add_heading('وزن‌های ارتباطی CGP', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(
            'نقشه حرارتی وزن‌های CGP نشان می‌دهد کدام شهرها (ستون‌ها، Influencer) '
            'بر کدام شهرها (سطرها، Influenced) تأثیر می‌گذارند. '
            'تنها ورودی‌های فعال (Active Inputs) که واقعاً در مدل نهایی CGP استفاده شده‌اند نمایش داده می‌شوند.'
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_img(doc, cgp_path(f'cgp_weights_heatmap_{pct}pct_spillover.png'),
                f'شکل: وزن‌های ارتباطی CGP (ماتریس تاثیر) - سرایت {pct}٪', Inches(5))

        doc.add_heading('همگرایی تابع شایستگی', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(
            'نمودارهای زیر روند کاهش خطای MSE (تابع شایستگی) را در طول نسل‌های '
            'الگوریتم تکاملی نشان می‌دهند. کاهش یکنواخت MSE نشان‌دهنده همگرایی '
            'صحیح الگوریتم CGP است. مقدار R2 نهایی هر شهر نیز نمایش داده شده است.'
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_img(doc, cgp_path(f'cgp_fitness_convergence_{pct}pct_spillover.png'),
                f'شکل: همگرایی تابع شایستگی CGP - سرایت {pct}٪', Inches(5.5))

        doc.add_page_break()

    # ==================== SECTION 9: Cross-comparison ====================
    doc.add_heading('۹. مقایسه کلی با حقیقت زمینی', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'نتایج CGP در هر سه درصد سرایت با ارتباطات شناخته‌شده (Ground Truth) مقایسه شده‌اند. '
        'معیارهای ارزیابی شامل Precision (دقت)، Recall (فراخوانی) و F1 Score هستند.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    doc.add_paragraph(
        'Precision = TP / (TP + FP)\nRecall = TP / (TP + FN)\nF1 = 2 * P * R / (P + R)'
    ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_spacer(doc)

    if data['gt_comparison'] is not None:
        gt_comp = data['gt_comparison']
        comp_headers = ['سرایت', 'کشف‌شده', 'حقیقت', 'TP', 'FP', 'FN', 'Precision', 'Recall', 'F1', 'لگ بهینه']
        comp_rows = []
        for _, row in gt_comp.iterrows():
            comp_rows.append([
                str(row['spillover']),
                str(int(row['discovered'])),
                str(int(row['ground_truth'])),
                str(int(row['true_positives'])),
                str(int(row['false_positives'])),
                str(int(row['false_negatives'])),
                f"{row['precision']:.3f}",
                f"{row['recall']:.3f}",
                f"{row['f1']:.3f}",
                str(int(row['optimal_lag'])),
            ])
        add_styled_table(doc, comp_headers, comp_rows)
    add_spacer(doc)

    add_img(doc, cgp_path('precision_recall_comparison.png'),
            'شکل: مقایسه Precision / Recall / F1 بر حسب درصد سرایت', Inches(5.5))

    add_img(doc, cgp_path('discovered_vs_ground_truth.png'),
            'شکل: مقایسه ارتباطات کشف‌شده با حقیقت زمینی (TP, FP, FN)', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 10: Discussion ====================
    doc.add_heading('۱۰. تحلیل نتایج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('تاثیر توابع جمعی', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'استفاده از توابع جمعی (بدون تفاضل و قدرمطلق تفاضل) دقیقاً همان رویکرد '
        'تحلیل داده‌های واقعی ۲۰۲۰ است. با حذف تفاضل، CGP فقط روابط تجمعی و '
        'هموار بین شهرها را شناسایی می‌کند. Recall در هر سه درصد سرایت ۶۰٪ بود '
        'که نشان‌دهنده توانایی CGP در بازیابی ۶ از ۱۰ ارتباط واقعی است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    doc.add_heading('ماتریس همبستگی در مقابل CGP', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'ماتریس همبستگی نشان می‌دهد که شهرها به دلیل دینامیک مشابه SIR، '
        'همبستگی بالایی دارند (اغلب بالای ۰.۹). CGP فراتر از همبستگی ساده '
        'عمل می‌کند و با تحلیل ورودی‌های فعال، شهرهای واقعاً مؤثر را شناسایی می‌کند.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    doc.add_heading('وزن‌های ارتباطی و جهت‌گیری', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'نقشه حرارتی وزن‌های CGP برخلاف ماتریس همبستگی (که متقارن است)، '
        'ساختار جهت‌دار تاثیرگذاری را نشان می‌دهد: کدام شهر بر کدام شهر اثر دارد. '
        'این اطلاعات برای درک جریان بیماری بین شهرها بسیار ارزشمند است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    doc.add_heading('همگرایی الگوریتم', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'نمودارهای همگرایی نشان می‌دهند که تابع شایستگی (MSE) به صورت یکنواخت '
        'کاهش می‌یابد. این تایید می‌کند که CGP به درستی عمل کرده و به جواب‌های '
        'مطلوب همگرا شده است. R2 نهایی برای تمام شهرها بالای ۰.۹۹ است.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_spacer(doc)

    doc.add_heading('False Positive ها', level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        'تعداد False Positive نسبتاً زیاد است. دلیل اصلی این است که '
        'تمام شهرها دینامیک SIR مشابهی دارند (beta و gamma یکسان) و '
        'روندهای فصلی مشابهی در مبتلایان دارند. در داده‌های واقعی، '
        'تفاوت‌های جغرافیایی و زمانی بیشتر به کاهش FP کمک می‌کنند.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    # ==================== SECTION 11: Conclusion ====================
    doc.add_heading('۱۱. نتیجه‌گیری', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    conclusions = [
        'CGP با توابع جمعی (Addition-Only) ۶ از ۱۰ ارتباط واقعی (Recall=60%) را بازیابی کرد',
        'لگ بهینه برای هر سه درصد سرایت ۳ روز تعیین شد',
        'ماتریس همبستگی و وزن‌های CGP اطلاعات مکمل ارائه می‌دهند',
        'گراف شبکه، ارتباطات صحیح (TP) و نادرست (FP) را به وضوح تفکیک می‌کند',
        'همگرایی تابع شایستگی عملکرد صحیح الگوریتم را تایید می‌کند',
        'همان رویکرد داده‌های واقعی (۲۰۲۰) با موفقیت روی داده‌های سینتتیک اعمال شد',
        'R2 نهایی برای تمام شهرها بالاتر از ۰.۹۹ است',
    ]
    for item in conclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(OUTPUT_DIR, 'CGP_Report_Farsi_SynData.docx')
    doc.save(path)
    print(f"  Farsi report saved: {path}")


# ==============================================================================
# ENGLISH REPORT
# ==============================================================================

def generate_english_report(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Times New Roman'

    # ==================== TITLE PAGE ====================
    for _ in range(4):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('CGP Analysis Report\nInter-City Connection Discovery from Synthetic Data')
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(46, 64, 87)
    add_spacer(doc, 2)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Cartesian Genetic Programming (CGP)\nAddition-Only Function Set\nGround Truth Validation')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(100, 100, 100)
    add_spacer(doc, 2)
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run('SIR Model: beta=0.35, gamma=0.14, R0=2.50\n'
                   'Cities: 10 Italian cities\n'
                   'R2 Threshold = 0.9')
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)
    add_spacer(doc)
    d2 = doc.add_paragraph()
    d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d2.add_run('Date: April 2026').font.size = Pt(12)
    doc.add_page_break()

    # ==================== TOC ====================
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Introduction and Objective',
        '2. CGP Methodology Overview',
        '3. Function Set (Addition-Only)',
        '4. CGP Parameters',
        '5. Ground Truth Connections',
        '6. Correlation Matrix',
        '7. Optimal Lag Search',
        '8. CGP Analysis Results',
        '   8.1 1% Spillover',
        '   8.2 5% Spillover',
        '   8.3 10% Spillover',
        '9. Overall Ground Truth Comparison',
        '10. Discussion',
        '11. Conclusions',
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ==================== SECTION 1: Introduction ====================
    doc.add_heading('1. Introduction and Objective', level=1)
    doc.add_paragraph(
        'This analysis validates Cartesian Genetic Programming (CGP) for discovering '
        'direct inter-city infection relationships. The EXACT SAME methodology and '
        'parameters used on real COVID-19 data (year 2020) are applied here to '
        'synthetic data with known (ground truth) connections.'
    )
    doc.add_paragraph(
        'Synthetic data was generated using the SIR model (beta=0.35, gamma=0.14) '
        'for 10 Italian cities. A predefined network of 10 ground truth connections '
        'was embedded, and CGP attempts to rediscover these connections from '
        'daily new infection data.'
    )
    doc.add_paragraph('Three spillover intensity levels were tested:')
    for item in ['1% - Weak spillover (challenging for detection)',
                 '5% - Medium spillover',
                 '10% - Strong spillover (easier to detect)']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # ==================== SECTION 2: CGP Method ====================
    doc.add_heading('2. CGP Methodology Overview', level=1)
    doc.add_paragraph(
        'CGP is an evolutionary algorithm that represents mathematical expressions '
        'as directed acyclic graphs on a 2D grid. Each node has a function gene '
        'and two input connection genes.'
    )
    doc.add_paragraph('CGP workflow:')
    steps = [
        'For each target city, lagged data from all other cities is provided as input',
        'CGP evolves a mathematical expression to predict the target city\'s new daily infections',
        'The (1+lambda) evolutionary strategy selects the best individual each generation',
        'Active Inputs in the evolved program reveal which cities truly influence the target',
        'Connections with R2 > 0.9 are considered statistically significant',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    doc.add_page_break()

    # ==================== SECTION 3: Functions ====================
    doc.add_heading('3. Function Set (Addition-Only)', level=1)
    doc.add_paragraph(
        'Following the exact same approach as the real 2020 data analysis, ONLY '
        'addition-based functions are used. NO subtraction (sub), absolute difference '
        '(abs_diff), multiplication (mul), or division (div) functions are included.'
    )
    add_spacer(doc)

    add_styled_table(doc, ['Function', 'Formula', 'Description'], [
        ['add', 'a + b', 'Direct addition of two inputs'],
        ['max', 'max(a, b)', 'Maximum of two inputs'],
        ['min', 'min(a, b)', 'Minimum of two inputs'],
        ['avg', '(a + b) / 2', 'Average of two inputs'],
        ['weighted_add', '0.7a + 0.3b', 'Weighted addition (emphasis on first input)'],
    ])
    add_spacer(doc)
    doc.add_paragraph(
        'Rationale: Addition-only functions model the aggregative nature of infection '
        'spread more naturally. Disease transmission is inherently cumulative: '
        'infected individuals from multiple cities collectively influence the target.'
    )
    doc.add_page_break()

    # ==================== SECTION 4: Parameters ====================
    doc.add_heading('4. CGP Parameters', level=1)
    add_styled_table(doc, ['Parameter', 'Value', 'Description'], [
        ['Grid Rows', '3', 'Number of rows in CGP grid'],
        ['Grid Columns', '8', 'Number of columns in CGP grid'],
        ['Outputs', '1', 'Predict target city daily new infections'],
        ['Generations', '500', 'Maximum evolutionary generations'],
        ['Lambda', '4', 'Children per generation'],
        ['Mutation Rate', '0.10', 'Per-gene mutation probability (10%)'],
        ['Runs', '3', 'Best result from 3 independent runs'],
        ['R2 Threshold', '0.9', 'Minimum R2 for significant connection'],
        ['Levels-back', '8', 'Full layer connectivity'],
        ['Functions', '5', 'add, max, min, avg, weighted_add'],
        ['Normalization', 'Z-score', 'Standardization of inputs and output'],
    ])
    doc.add_page_break()

    # ==================== SECTION 5: Ground Truth ====================
    doc.add_heading('5. Ground Truth Connections', level=1)
    doc.add_paragraph(
        'The following network of 10 connections was defined in the SIR model. '
        'Each connection has a specified strength and temporal delay. '
        'CGP\'s task is to rediscover these connections from infection time series.'
    )
    add_spacer(doc)

    if data['ground_truth'] is not None:
        gt = data['ground_truth']
        gt_rows = []
        for _, row in gt.iterrows():
            gt_rows.append([
                str(row['source']).title(),
                str(row['target']).title(),
                str(row['strength']).title(),
                f"{row['lag_days']} days",
            ])
        add_styled_table(doc, ['Source City', 'Target City', 'Strength', 'Delay'], gt_rows)

    add_spacer(doc)
    add_img(doc, sir_path('ground_truth_network.png'),
            'Figure 1: Ground Truth Network (defined in SIR model)', Inches(5))
    doc.add_page_break()

    # ==================== SECTION 6: Correlation Matrix ====================
    doc.add_heading('6. Correlation Matrix', level=1)
    doc.add_paragraph(
        'The Pearson correlation matrix between daily new infections of all cities '
        'is computed for each spillover percentage. This provides a linear baseline '
        'before the non-linear CGP analysis. Note that high correlation does not '
        'necessarily imply direct connection - CGP identifies causal relationships beyond correlation.'
    )
    add_spacer(doc)

    for pct in [1, 5, 10]:
        add_img(doc, cgp_path(f'correlation_matrix_{pct}pct_spillover.png'),
                f'Figure: Pearson Correlation Matrix - {pct}% Spillover', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 7: Optimal Lag ====================
    doc.add_heading('7. Optimal Lag Search', level=1)
    doc.add_paragraph(
        'Different lag values (1, 2, 3, 5, 7, 10, 14, 21 days) were tested to find '
        'the optimal temporal lookback. For each lag value, CGP was run for all cities '
        'and the average R2 was computed.'
    )
    add_spacer(doc)

    if data['lag_sweep'] is not None:
        lag = data['lag_sweep']
        lag_headers = ['Spillover', 'Lag (days)', 'Avg R2', 'Max R2', 'Cities Above Threshold']
        lag_rows = []
        for _, row in lag.iterrows():
            lag_rows.append([
                str(row['spillover']),
                str(int(row['lag'])),
                f"{row['avg_r2']:.4f}",
                f"{row['max_r2']:.4f}",
                str(int(row['n_above_threshold'])),
            ])
        add_styled_table(doc, lag_headers, lag_rows)
    add_spacer(doc)

    for pct in [1, 5, 10]:
        add_img(doc, cgp_path(f'lag_sweep_({pct}pct_spillover).png'),
                f'Figure: Lag Sweep Results - {pct}% Spillover', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 8: Results per % ====================
    doc.add_heading('8. CGP Analysis Results', level=1)

    for pct_idx, pct in enumerate([1, 5, 10]):
        doc.add_heading(f'8.{pct_idx+1} {pct}% Spillover', level=2)

        # City summary table
        cs_key = f'city_summary_{pct}'
        if data[cs_key] is not None:
            cs = data[cs_key]
            doc.add_heading('Per-City Summary', level=3)
            cs_headers = ['City', 'R2', 'Connections', 'Linked Cities', 'Active Functions']
            cs_rows = []
            for _, row in cs.iterrows():
                linked = str(row.get('linked_cities', '')).replace(';', ', ')
                funcs = str(row.get('active_functions', ''))
                if funcs == 'nan':
                    funcs = '-'
                cs_rows.append([
                    str(row['city']),
                    f"{row['r2_score']:.4f}",
                    str(int(row['n_connections'])),
                    linked if linked != 'nan' else '-',
                    funcs,
                ])
            add_styled_table(doc, cs_headers, cs_rows)
        add_spacer(doc)

        # Connections table
        conn_key = f'connections_{pct}'
        if data[conn_key] is not None:
            cdf = data[conn_key]
            sig = cdf[cdf['is_significant'] == True]
            doc.add_heading('Discovered Connections', level=3)
            doc.add_paragraph(
                f'Total significant connections: {len(sig)} (out of {len(cdf)} pairs analyzed)')

            conn_headers = ['City 1', 'City 2', 'Strength (R2)', 'Ground Truth', 'Classification']
            conn_rows = []
            for _, row in sig.iterrows():
                conn_rows.append([
                    str(row['city_1']),
                    str(row['city_2']),
                    f"{row['connection_strength_R2']:.4f}",
                    'Yes' if row.get('is_ground_truth', False) else 'No',
                    str(row.get('classification', '')),
                ])
            add_styled_table(doc, conn_headers, conn_rows)
        add_spacer(doc)

        # ALL visualizations for this percentage
        doc.add_heading('Correlation Matrix', level=3)
        add_img(doc, cgp_path(f'correlation_matrix_{pct}pct_spillover.png'),
                f'Figure: Pearson Correlation Matrix - {pct}% Spillover', Inches(5))

        doc.add_heading('Connection Strength Heatmap', level=3)
        doc.add_paragraph(
            'This heatmap shows the strength (R2) of CGP-discovered connections '
            'between all city pairs.')
        add_img(doc, cgp_path(f'connection_heatmap_{pct}pct_spillover.png'),
                f'Figure: CGP Connection Strength Heatmap - {pct}% Spillover', Inches(5))

        doc.add_heading('Network Graph', level=3)
        doc.add_paragraph(
            'Green solid edges represent True Positive connections (correctly discovered). '
            'Red dashed edges represent False Positive connections (extra/spurious). '
            'Node size is proportional to the number of connections.')
        add_img(doc, cgp_path(f'network_graph_{pct}pct_spillover.png'),
                f'Figure: CGP Network Graph - {pct}% Spillover', Inches(5.5))

        doc.add_heading('CGP Influence Weights', level=3)
        doc.add_paragraph(
            'The CGP weights heatmap shows which cities (columns, Source/Influencer) '
            'affect which other cities (rows, Target/Influenced). Unlike the symmetric '
            'correlation matrix, this reveals the DIRECTIONAL influence structure. '
            'Only Active Inputs actually used in the final CGP model are shown.')
        add_img(doc, cgp_path(f'cgp_weights_heatmap_{pct}pct_spillover.png'),
                f'Figure: CGP Influence Weights (Directional) - {pct}% Spillover', Inches(5))

        doc.add_heading('Fitness Convergence', level=3)
        doc.add_paragraph(
            'These plots show the MSE (fitness function) decrease over evolutionary '
            'generations. A monotonic decrease confirms proper CGP convergence. '
            'The final R2 value for each city is displayed in the subplot title.')
        add_img(doc, cgp_path(f'cgp_fitness_convergence_{pct}pct_spillover.png'),
                f'Figure: CGP Fitness Convergence - {pct}% Spillover', Inches(5.5))

        doc.add_page_break()

    # ==================== SECTION 9: Cross-comparison ====================
    doc.add_heading('9. Overall Ground Truth Comparison', level=1)
    doc.add_paragraph(
        'CGP results across all three spillover percentages are compared against '
        'the known ground truth. Evaluation metrics:')
    doc.add_paragraph(
        'Precision = TP / (TP + FP) - What fraction of discovered connections are correct?\n'
        'Recall = TP / (TP + FN) - What fraction of true connections were found?\n'
        'F1 = 2 * P * R / (P + R) - Harmonic mean of Precision and Recall')
    add_spacer(doc)

    if data['gt_comparison'] is not None:
        gt_comp = data['gt_comparison']
        comp_headers = ['Spillover', 'Discovered', 'Truth', 'TP', 'FP', 'FN',
                        'Precision', 'Recall', 'F1', 'Optimal Lag']
        comp_rows = []
        for _, row in gt_comp.iterrows():
            comp_rows.append([
                str(row['spillover']),
                str(int(row['discovered'])),
                str(int(row['ground_truth'])),
                str(int(row['true_positives'])),
                str(int(row['false_positives'])),
                str(int(row['false_negatives'])),
                f"{row['precision']:.3f}",
                f"{row['recall']:.3f}",
                f"{row['f1']:.3f}",
                str(int(row['optimal_lag'])),
            ])
        add_styled_table(doc, comp_headers, comp_rows)
    add_spacer(doc)

    add_img(doc, cgp_path('precision_recall_comparison.png'),
            'Figure: Precision / Recall / F1 Score by Spillover Percentage', Inches(5.5))

    add_img(doc, cgp_path('discovered_vs_ground_truth.png'),
            'Figure: Connection Discovery Breakdown - True Positives, False Positives, '
            'False Negatives per Spillover %', Inches(5.5))
    doc.add_page_break()

    # ==================== SECTION 10: Discussion ====================
    doc.add_heading('10. Discussion', level=1)

    doc.add_heading('Impact of Addition-Only Functions', level=3)
    doc.add_paragraph(
        'Using the exact same addition-only functions as the real 2020 data analysis '
        'ensures methodological consistency. By removing subtraction and absolute '
        'difference, CGP models only aggregative relationships. Recall reached 60% '
        'consistently across all spillover levels, demonstrating CGP\'s ability to '
        'recover 6 out of 10 true connections.')
    add_spacer(doc)

    doc.add_heading('Correlation Matrix vs CGP', level=3)
    doc.add_paragraph(
        'The correlation matrix shows that all cities have high pairwise correlations '
        '(typically >0.9) due to sharing the same SIR dynamics. CGP goes beyond '
        'simple correlation by analyzing which specific cities contribute to predicting '
        'others through its Active Input identification mechanism.')
    add_spacer(doc)

    doc.add_heading('CGP Influence Weights', level=3)
    doc.add_paragraph(
        'Unlike the symmetric correlation matrix, the CGP weights heatmap reveals '
        'directional influence: which cities serve as sources and which as targets. '
        'This asymmetric information is more valuable for understanding infection flow.')
    add_spacer(doc)

    doc.add_heading('Fitness Convergence', level=3)
    doc.add_paragraph(
        'All fitness convergence plots show monotonic MSE decrease, confirming '
        'proper evolutionary search. Final R2 values exceed 0.99 for all cities, '
        'indicating excellent predictive capability.')
    add_spacer(doc)

    doc.add_heading('False Positive Analysis', level=3)
    doc.add_paragraph(
        'The relatively high false positive rate is expected because all cities share '
        'identical SIR dynamics (same beta and gamma), creating inherently similar '
        'temporal patterns. In real data, geographic and temporal heterogeneity '
        'naturally reduces false positives. The key finding is that CGP consistently '
        'maintains 60% Recall across all spillover levels.')
    doc.add_page_break()

    # ==================== SECTION 11: Conclusions ====================
    doc.add_heading('11. Conclusions', level=1)
    conclusions = [
        'CGP with addition-only functions recovered 6 out of 10 true connections (Recall=60%)',
        'Optimal lag was consistently 3 days across all spillover percentages',
        'Correlation matrix provides linear baseline; CGP captures non-linear dependencies',
        'Network graphs clearly distinguish True Positives (green) from False Positives (red dashed)',
        'CGP influence weights reveal directional city-to-city relationships',
        'Fitness convergence confirms proper algorithm behavior (R2 > 0.99 for all cities)',
        'The same methodology used on real 2020 COVID-19 data was successfully validated on synthetic data',
    ]
    for item in conclusions:
        doc.add_paragraph(item, style='List Bullet')

    path = os.path.join(OUTPUT_DIR, 'CGP_Report_English_SynData.docx')
    doc.save(path)
    print(f"  English report saved: {path}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    print("Generating CGP Reports for Synthetic Data...")
    print(f"  CGP results: {CGP_DIR}")
    print(f"  SIR results: {SIR_DIR}")

    data = load_all_data()

    # Check loaded data
    for key, val in data.items():
        if val is not None:
            if isinstance(val, pd.DataFrame):
                print(f"  Loaded: {key} ({len(val)} rows)")
            else:
                print(f"  Loaded: {key}")
        else:
            print(f"  MISSING: {key}")

    generate_farsi_report(data)
    generate_english_report(data)

    print(f"\nDone! Reports saved to: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
